"""Build V8 patent disclosure with direct flowchart feedback line."""

from __future__ import annotations

import importlib.util
import re
import subprocess
from datetime import datetime
from pathlib import Path

from PIL import Image, ImageDraw


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = Path(__file__).resolve().parent
FIG_DIR = OUT_DIR / "figures"
TMP_DIR = OUT_DIR / "_tmp"
V3_DIR = ROOT / "专利撰写" / "可微控制参数整定交底书_V3"
V3_FIG_DIR = V3_DIR / "figures"
MLP_DIAG_FIG = ROOT / "sim" / "results" / "diagnostic" / "mlp_instability" / "0507" / "ROOT_CAUSE_STORY.png"
CASE_NAME = "一种基于车辆动力学可微闭环仿真的车辆横纵向控制器参数自动整定方法"


def load_v3_module():
    spec = importlib.util.spec_from_file_location("build_disclosure_v3", V3_DIR / "build_disclosure_v3.py")
    if spec is None or spec.loader is None:
        raise RuntimeError("failed to load V3 builder")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


V3 = load_v3_module()


def arrowhead(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int]) -> None:
    sx, sy = start
    ex, ey = end
    if abs(ex - sx) >= abs(ey - sy):
        sign = 1 if ex >= sx else -1
        pts = [(ex, ey), (ex - sign * 14, ey - 8), (ex - sign * 14, ey + 8)]
    else:
        sign = 1 if ey >= sy else -1
        pts = [(ex, ey), (ex - 8, ey - sign * 14), (ex + 8, ey - sign * 14)]
    draw.polygon(pts, fill="black")


def poly_arrow(draw: ImageDraw.ImageDraw, points: list[tuple[int, int]]) -> None:
    draw.line(points, fill="black", width=2)
    arrowhead(draw, points[-2], points[-1])


def straight_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int]) -> None:
    poly_arrow(draw, [start, end])


def center(box: tuple[int, int, int, int]) -> tuple[int, int]:
    x1, y1, x2, y2 = box
    return ((x1 + x2) // 2, (y1 + y2) // 2)


def generate_system_diagram(path: Path) -> None:
    img = Image.new("RGB", (1800, 1000), "white")
    draw = ImageDraw.Draw(img)
    draw.text((60, 35), "图1  可微闭环整定系统框图", fill="black", font=V3.font(34, bold=True))

    boxes = {
        "A": (70, 330, 350, 450, "原始控制器\n代码/参数表"),
        "B": (430, 330, 710, 450, "双模式控制器\n训练路径/验证路径"),
        "C": (790, 330, 1080, 450, "闭环展开模块\n横向-纵向-动力学"),
        "D": (1160, 330, 1450, 450, "综合损失模块\n误差/平滑/正则"),
        "E": (1510, 330, 1740, 450, "自动微分模块\n梯度回传"),
        "F": (690, 110, 940, 230, "混合被控对象\n机理模型+MLP残差"),
        "G": (970, 110, 1220, 230, "多场景轨迹库\n轨迹/速度段"),
        "H": (1240, 110, 1490, 230, "训练增强模块\n域随机化/噪声/抖动"),
        "I": (1160, 620, 1450, 750, "参数投影模块\n物理边界约束"),
        "J": (790, 620, 1080, 750, "产物生成模块\n整定配置/日志"),
        "K": (430, 620, 710, 750, "硬逻辑复验模块\n原始分支/限幅/速率"),
        "L": (1510, 620, 1740, 750, "MLP诊断模块\n输出/OOD/消融"),
    }
    highlight = {"B", "C", "D", "E", "I"}
    for key, (x1, y1, x2, y2, text) in boxes.items():
        V3.draw_box(draw, (x1, y1, x2, y2), text, fill="#F7F7F7" if key in highlight else "white")

    # Main forward chain.
    straight_arrow(draw, (350, 390), (430, 390))
    straight_arrow(draw, (710, 390), (790, 390))
    straight_arrow(draw, (1080, 390), (1160, 390))
    straight_arrow(draw, (1450, 390), (1510, 390))

    # Orthogonal input connections into the closed-loop and loss path.
    straight_arrow(draw, (815, 230), (815, 330))
    poly_arrow(draw, [(1095, 230), (1095, 280), (1015, 280), (1015, 330)])
    straight_arrow(draw, (1365, 230), (1365, 330))

    # Parameter update, output, validation, and outer feedback loop.
    poly_arrow(draw, [(1625, 450), (1625, 535), (1305, 535), (1305, 620)])
    straight_arrow(draw, (1160, 685), (1080, 685))
    straight_arrow(draw, (790, 685), (710, 685))
    poly_arrow(draw, [(1450, 685), (1510, 685)])
    poly_arrow(draw, [(1625, 750), (1625, 875), (935, 875), (935, 750)])
    poly_arrow(draw, [(570, 750), (570, 870), (745, 870), (745, 535), (935, 535), (935, 450)])
    draw.text((575, 835), "复验未达标则继续整定", fill="black", font=V3.font(20))

    draw.text(
        (70, 930),
        "说明：混合被控对象用于提高车辆响应逼真度，诊断模块用于区分控制器问题与模型残差失效。",
        fill="black",
        font=V3.font(18),
    )
    img.save(path)


def generate_flow_diagram(path: Path) -> None:
    img = Image.new("RGB", (1800, 1050), "white")
    draw = ImageDraw.Draw(img)
    draw.text((60, 35), "图2  控制器参数自动整定流程图", fill="black", font=V3.font(34, bold=True))

    w, h = 310, 125
    y_top, y_bottom = 155, 450
    xs = [100, 525, 950, 1375]
    steps = {
        "S1": (xs[0], y_top, xs[0] + w, y_top + h, "S1\n读取控制器逻辑\n识别参数与边界"),
        "S2": (xs[1], y_top, xs[1] + w, y_top + h, "S2\n构造双模式控制器\n可微训练/硬逻辑验证"),
        "S3": (xs[2], y_top, xs[2] + w, y_top + h, "S3\n构造混合被控对象\n机理+MLP残差"),
        "S4": (xs[3], y_top, xs[3] + w, y_top + h, "S4\n构造多轨迹多速度\n批量闭环样本"),
        "S5": (xs[3], y_bottom, xs[3] + w, y_bottom + h, "S5\n按控制周期展开\n横向-纵向-动力学"),
        "S6": (xs[2], y_bottom, xs[2] + w, y_bottom + h, "S6\n计算综合损失\n误差/平滑/正则"),
        "S7": (xs[1], y_bottom, xs[1] + w, y_bottom + h, "S7\n反向传播更新\n执行物理投影"),
        "S8": (xs[0], y_bottom, xs[0] + w, y_bottom + h, "S8\n硬逻辑复验\n导出整定配置"),
        "S9": (100, 735, 410, 860, "S9\nMLP可视化诊断\n定位模型或控制器因素"),
    }
    for key, (x1, y1, x2, y2, text) in steps.items():
        V3.draw_box(draw, (x1, y1, x2, y2), text, title=key in {"S1", "S8"}, fill="#F7F7F7")

    straight_arrow(draw, (steps["S1"][2], center(steps["S1"][:4])[1]), (steps["S2"][0], center(steps["S2"][:4])[1]))
    straight_arrow(draw, (steps["S2"][2], center(steps["S2"][:4])[1]), (steps["S3"][0], center(steps["S3"][:4])[1]))
    straight_arrow(draw, (steps["S3"][2], center(steps["S3"][:4])[1]), (steps["S4"][0], center(steps["S4"][:4])[1]))
    straight_arrow(draw, (center(steps["S4"][:4])[0], steps["S4"][3]), (center(steps["S5"][:4])[0], steps["S5"][1]))
    straight_arrow(draw, (steps["S5"][0], center(steps["S5"][:4])[1]), (steps["S6"][2], center(steps["S6"][:4])[1]))
    straight_arrow(draw, (steps["S6"][0], center(steps["S6"][:4])[1]), (steps["S7"][2], center(steps["S7"][:4])[1]))
    straight_arrow(draw, (steps["S7"][0], center(steps["S7"][:4])[1]), (steps["S8"][2], center(steps["S8"][:4])[1]))
    straight_arrow(draw, (center(steps["S8"][:4])[0], steps["S8"][3]), (center(steps["S9"][:4])[0], steps["S9"][1]))

    poly_arrow(
        draw,
        [
            (steps["S9"][2], center(steps["S9"][:4])[1]),
            (center(steps["S5"][:4])[0], center(steps["S9"][:4])[1]),
            (center(steps["S5"][:4])[0], steps["S5"][3]),
        ],
    )
    label = "控制器不足则继续整定；模型异常则检查 MLP 残差"
    label_font = V3.font(20)
    label_bbox = draw.textbbox((0, 0), label, font=label_font)
    label_x = steps["S9"][2] + (center(steps["S5"][:4])[0] - steps["S9"][2] - (label_bbox[2] - label_bbox[0])) // 2
    draw.text((label_x, center(steps["S9"][:4])[1] + 32), label, fill="black", font=label_font)
    img.save(path)


def generate_figures() -> None:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    generate_system_diagram(FIG_DIR / "fig1_system_architecture.png")
    generate_flow_diagram(FIG_DIR / "fig2_method_flow.png")
    for name in [
        "fig3_loss_curve.png",
        "fig4_parameter_changes.png",
        "fig5_training_summary.png",
        "fig6_comparison_trajectory.png",
        "fig7_comparison_lateral_error.png",
    ]:
        V3.grayscale_copy(V3_FIG_DIR / name, FIG_DIR / name)
    V3.grayscale_copy(MLP_DIAG_FIG, FIG_DIR / "fig8_mlp_diagnostic_story.png")


SYSTEM_MERMAID = r"""
```mermaid
flowchart LR
  A["原始控制器代码/参数表"] --> B["双模式控制器"]
  B --> C["闭环展开模块"]
  D["混合被控对象：机理模型+MLP残差"] --> C
  E["多场景轨迹库"] --> C
  F["训练增强模块：域随机化/噪声/抖动"] --> C
  C --> G["综合损失模块"]
  G --> H["自动微分模块"]
  H --> I["参数投影模块"]
  I --> J["产物生成模块"]
  J --> K["硬逻辑复验模块"]
  J --> L["MLP可视化诊断模块"]
  K --"未达标继续整定"--> C
  L --"模型残差异常"--> D
```
<!-- ![图1 可微闭环整定系统框图](figures/fig1_system_architecture.png) -->
"""


FLOW_MERMAID = r"""
```mermaid
flowchart TB
  S1["S1 读取控制器逻辑并识别参数"] --> S2["S2 构造双模式控制器"]
  S2 --> S3["S3 构造机理+MLP残差被控对象"]
  S3 --> S4["S4 构造多轨迹多速度训练集"]
  S4 --> S5["S5 按控制周期闭环展开"]
  S5 --> S6["S6 计算综合损失"]
  S6 --> S7["S7 反向传播并投影更新参数"]
  S7 --> S8["S8 硬逻辑复验并导出配置"]
  S8 --> S9["S9 MLP可视化诊断"]
  S8 --"不满足验收"--> S5
  S9 --"模型残差异常"--> S3
```
<!-- ![图2 控制器参数自动整定流程](figures/fig2_method_flow.png) -->
"""


def replace_section(md: str, start: str, end: str, replacement: str) -> str:
    pattern = re.escape(start) + r".*?(?=" + re.escape(end) + r")"
    updated, count = re.subn(pattern, lambda _m: replacement, md, flags=re.S)
    if count != 1:
        raise RuntimeError(f"expected to replace one section starting with {start!r}, got {count}")
    return updated


def formalize_markdown(md: str) -> str:
    md = re.sub(r"\n\*\*版本说明\*\*：.*?\n", "\n", md)
    md = re.sub(r"\n---\n\n## 注意事项\n.*?\n## 一、", "\n---\n\n## 一、", md, flags=re.S)

    replacements = {
        "## 一、介绍相关技术背景，描述与本发明技术最相近的现有技术，并说明该现有技术存在的缺点": "## 一、技术背景、最接近现有技术及现有技术缺点",
        "## 二、针对上述缺点，说明本发明所要解决的技术问题": "## 二、本发明所要解决的技术问题",
        "## 四、与现有技术相比，本发明具有哪些优点？": "## 四、与现有技术相比的有益效果",
        "## 五、本发明的技术关键点和欲保护点是什么？": "## 五、本发明的技术关键点和保护点",
        "## 六、其它（实施例、技术效果、参数示例）": "## 六、实施例、技术效果和参数示例",
    }
    for old, new in replacements.items():
        md = md.replace(old, new)

    md = md.replace(V3.SYSTEM_MERMAID.strip(), SYSTEM_MERMAID.strip())
    md = md.replace(V3.FLOW_MERMAID.strip(), FLOW_MERMAID.strip())
    md = md.replace("可部署 YAML 参数配置", "可部署参数配置")
    md = md.replace(
        "图6和图7示出了代表场景的轨迹跟踪和横向误差对比；完整结果保留在项目结果目录中。",
        "图6和图7示出了代表场景的轨迹跟踪和横向误差对比；完整验证材料可随整定日志一并保存。",
    )
    md = md.replace(" kph", " km/h")

    section2 = """## 二、本发明所要解决的技术问题

本发明所要解决的技术问题在于，在不改变既有车辆横向、纵向工程控制器主体结构的前提下，构建能够用于参数自动整定的车辆动力学可微闭环仿真环境，使控制器参数能够根据轨迹跟踪误差、速度误差和控制平滑性指标自动更新。

进一步地，本发明还解决训练可微性与工程可部署性之间的冲突：训练阶段允许对限幅、分段切换、速率限制等非光滑环节采用可微近似，以便获得参数更新方向；验证和部署前仍使用原始硬限幅、硬分支和硬速率限制进行闭环复验，以保证整定参数能够回到原工程控制模块中使用。

此外，本发明还解决被控对象精度和整定可微性之间的协调问题：车辆运动响应由机理动力学模型给出主体趋势，再由 MLP 残差模型补偿机理模型与高保真车辆响应之间的差异。MLP 残差模型的权重在控制器参数整定阶段保持冻结，系统仅利用其输入输出关系修正车辆状态并传递梯度，从而使控制器参数在更接近真实车辆响应的闭环环境中自动更新。

进一步地，本发明还解决闭环失控原因难以归因的问题。系统在训练或验证后捕获 MLP 输入、归一化距离、残差输出和车辆响应，形成开环静态扫描、闭环时序、输入分布外距离、组件消融和跨场景汇总等可视化诊断结果，用于判断异常主要来源于控制器参数不足、车辆机理模型失配，还是 MLP 残差模型在分布外输入下失效。

本发明还解决多轨迹、多速度和车辆物理参数不确定性下的统一整定问题，使横向预瞄、收敛、角速度误差预瞄、纵向位置环和速度环等多类参数能够在同一批量闭环训练过程中协同更新，并通过物理边界约束、硬逻辑复验、诊断图和日志产物保证整定过程可追溯。

"""
    md = replace_section(md, "## 二、本发明所要解决的技术问题", "## 三、本发明技术方案的详细阐述", section2)

    md = md.replace(
        "## 四、与现有技术相比的有益效果\n\n（1）",
        "## 四、与现有技术相比的有益效果\n\n本发明相较于现有技术至少具有以下有益效果。\n\n（1）",
    )
    md = md.replace(
        "## 五、本发明的技术关键点和保护点\n\n（1）",
        "## 五、本发明的技术关键点和保护点\n\n本发明建议重点保护以下技术方案和技术特征。\n\n（1）",
    )

    section64 = """### 6.4 可实施性说明

上述实施例可以采用通用计算设备、仿真服务器或车端开发环境实施。实施时，将车辆横向控制逻辑、纵向控制逻辑、车辆动力学模型、训练场景库、参数边界表和验证场景库配置在同一整定系统中；训练完成后，系统输出与原控制器参数格式一致或可转换为原参数格式的整定配置。

在一个可选实施方式中，系统保存每轮训练的参数变化、损失变化、代表场景跟踪曲线、硬逻辑复验统计和配置导出日志。上述材料用于说明参数整定过程、复验结果和部署前差异，不限定本发明的具体软件目录、编程语言、配置文件名称或图表样式。
"""
    if "### 6.4 项目实现依据" in md:
        md = md[: md.index("### 6.4 项目实现依据")] + section64
    md = merge_v6_content(md)
    return add_equation_numbers(md).strip() + "\n"


def merge_v6_content(md: str) -> str:
    md = md.replace(
        "（6）现有方法缺乏训练路径与验证路径分离而参数共享的结构，难以兼顾“可微训练需要平滑近似”和“工程部署需要原始硬逻辑”两类要求。",
        "（6）现有方法缺乏训练路径与验证路径分离而参数共享的结构，难以兼顾“可微训练需要平滑近似”和“工程部署需要原始硬逻辑”两类要求。\n\n"
        "（7）现有车辆控制器整定方案通常把被控对象简化为纯机理模型或固定仿真器，较少公开将机理动力学模型与冻结的 MLP 残差模型组合为可微被控对象，并进一步通过 MLP 输入输出可视化、分布外距离和组件消融来区分控制器问题与车辆模型失效。"
    )
    md = md.replace(
        "本发明将已有工程控制器复现为双模式结构：训练模式保留主要物理关系并对非光滑环节采用可微近似，验证模式保留原始硬逻辑。两种模式共享同一组待整定参数。通过车辆动力学模型把多个控制周期串联成闭环时间链路，综合损失可以沿时间反向传播到控制器参数，从而自动得到参数更新方向。",
        "本发明将已有工程控制器复现为双模式结构：训练模式保留主要物理关系并对非光滑环节采用可微近似，验证模式保留原始硬逻辑。两种模式共享同一组待整定参数。通过车辆动力学模型把多个控制周期串联成闭环时间链路，综合损失可以沿时间反向传播到控制器参数，从而自动得到参数更新方向。\n\n"
        "在重卡或牵引车-挂车实施场景中，被控对象可以采用机理动力学模型与 MLP 残差模型组合的形式。机理模型根据车辆质量、轴距、侧偏刚度、轮胎半径、传动扭矩和铰接约束等给出下一周期的名义状态；MLP 残差模型接收车辆状态、控制指令和车辆配置等特征，输出速度或位姿残差，并将该残差转换为车辆状态修正。该 MLP 残差模型作为被控对象的一部分使用，在控制器参数整定阶段不更新其权重。"
    )
    md = md.replace(
        "图1示出了本发明的一种系统组成。系统包括原始控制器代码/参数表、双模式控制器、车辆动力学模型、多场景轨迹库、闭环展开模块、损失计算模块、自动微分模块、参数投影模块、产物生成模块和硬逻辑复验模块。各模块之间形成从工程控制器复现、闭环仿真、参数更新到部署验证的闭环流程。",
        "图1示出了本发明的一种系统组成。系统包括原始控制器代码/参数表、双模式控制器、混合被控对象、多场景轨迹库、训练增强模块、闭环展开模块、损失计算模块、自动微分模块、参数投影模块、产物生成模块、硬逻辑复验模块和 MLP 可视化诊断模块。各模块之间形成从工程控制器复现、混合车辆响应仿真、参数更新、部署验证到异常归因诊断的闭环流程。"
    )

    section33 = """### 3.3 模块功能说明

（1）原始控制器代码/参数表用于提供待复现的横向和纵向工程控制逻辑，并给出初始标定参数、车辆物理常数和安全边界。

（2）双模式控制器用于在训练模式下提供可微计算路径，在验证模式下提供与工程控制器一致的硬逻辑路径。同一组整定参数在两个模式中共享。

（3）混合被控对象用于模拟车辆在控制指令作用下的下一周期响应。该被控对象至少包括机理动力学模型，并可进一步包括 MLP 残差模型。机理动力学模型负责表达车辆质量、轴距、轮胎侧偏、传动扭矩、铰接约束、空气阻力和滚阻等确定性物理关系；MLP 残差模型负责补偿机理模型与高保真仿真或实车数据之间的剩余误差。

（4）MLP 残差模型用于根据车辆状态、控制指令、车辆配置、铰接状态和归一化统计量等输入特征输出运动残差。运动残差可以包括牵引车速度残差、挂车速度残差和相对位姿残差。残差经坐标转换、限幅和必要的掩码处理后叠加到机理模型的下一周期状态上。MLP 权重在控制器整定阶段保持冻结，梯度仅通过其输入输出计算链路回传到控制器参数。

（5）多场景轨迹库用于提供换道、双换道、渐变曲率弯道、S弯、弯前减速和换道加减速等场景，并覆盖多个速度段。

（6）训练增强模块用于在训练阶段引入车辆物理参数域随机化、反馈噪声和指令抖动，使控制器参数不只适配单一车辆或单一干净输入。

（7）闭环展开模块用于按控制周期串联横向控制器、纵向控制器和混合被控对象，使每一时刻的状态变化都依赖前一时刻的控制结果。

（8）损失计算模块用于评价横向误差、航向误差、速度误差、转向平滑性、加速度平滑性和参数偏移，得到可优化的综合损失。

（9）自动微分模块用于沿闭环时间链路反向传播，计算综合损失对控制器参数的梯度。对于冻结的 MLP 残差模型，该模块不更新 MLP 权重，只利用其可微映射把车辆响应对控制器参数的影响传递回来。

（10）参数投影模块用于将更新后的参数限制在物理合理范围和安全范围内，防止自动整定越过工程可部署边界。

（11）产物生成模块用于输出整定后的参数配置、训练曲线、参数变化图、验证统计图、诊断图和日志。

（12）硬逻辑复验模块用于使用验证模式运行全场景闭环仿真，检查整定参数在原始硬分支、硬限幅和硬速率限制下的表现。

（13）MLP 可视化诊断模块用于在闭环异常或部署前复验不达标时，捕获 MLP 的输入特征、归一化后距离、原始残差输出、限幅后残差输出和车辆响应，并生成开环静态扫描、闭环时序图、输入分布外距离图、危险区热图、组件消融对比图和跨场景汇总图。通过比较“纯机理模型”“机理模型+完整 MLP”“屏蔽部分 MLP 输出分量”等结果，可以判断异常主要来自控制器整定不足，还是来自 MLP 残差模型输出偏置、饱和或分布外失效。

"""
    md = replace_section(md, "### 3.3 模块功能说明", "### 3.4 系统流程说明", section33)

    section34 = """### 3.4 系统流程说明

图2示出了本发明的一种方法流程。

```mermaid
flowchart TB
  S1["S1 读取控制器逻辑并识别参数"] --> S2["S2 构造双模式控制器"]
  S2 --> S3["S3 构造机理+MLP残差被控对象"]
  S3 --> S4["S4 构造多轨迹多速度训练集"]
  S4 --> S5["S5 按控制周期闭环展开"]
  S5 --> S6["S6 计算综合损失"]
  S6 --> S7["S7 反向传播并投影更新参数"]
  S7 --> S8["S8 硬逻辑复验并导出配置"]
  S8 --> S9["S9 MLP可视化诊断"]
  S8 --"不满足验收"--> S5
  S9 --"模型残差异常"--> S3
```
<!-- ![图2 控制器参数自动整定流程](figures/fig2_method_flow.png) -->

具体流程如下：

S1，读取车辆控制器的原始代码和参数配置，识别可调参数、固定物理参数和安全约束参数。

S2，将横向控制器和纵向控制器封装为双模式控制器。训练模式保留控制器主要物理关系，并对非光滑步骤采用可微近似；验证模式保留原始工程硬逻辑。

S3，构建机理模型与 MLP 残差模型组合的混合被控对象，并根据需要构建车辆物理参数域。机理模型可以采用运动学自行车模型、动力学自行车模型或牵引车-挂车双体动力学模型；MLP 残差模型用于修正机理模型的下一周期状态。

S4，构造多轨迹多速度训练集，并将轨迹场景、速度段和车辆参数域展开为批量样本。

S5，在每一个控制周期内，按横向控制、纵向控制、机理动力学预测和 MLP 残差修正的顺序推进闭环状态。

S6，根据车辆状态和参考轨迹计算综合损失。

S7，沿时间展开链路反向传播，获得损失对控制器参数的梯度，更新参数并执行物理约束投影；MLP 残差模型的权重保持冻结。

S8，导出整定后的参数配置，并在验证模式下进行硬逻辑复验；若复验不满足要求，则以上一轮参数为起点继续整定或调整训练场景。

S9，当硬逻辑复验出现异常、某些场景误差突然增大，或需要判断被控对象可信度时，运行 MLP 可视化诊断。若诊断显示 MLP 在无激励输入下产生持续偏置、输入远离训练分布、输出频繁触碰限幅或某一残差分量被屏蔽后异常消失，则优先判断为车辆模型残差失效；若纯机理路径和 MLP 路径均表现相似，则优先判断为控制器参数或控制器结构不足。

"""
    md = replace_section(md, "### 3.4 系统流程说明", "### 3.4.1 符号与公式", section34)

    section341 = r"""### 3.4.1 符号与公式

#### （1）符号与变量定义

| 符号 | 含义 | 下标/量纲 |
|------|------|-----------|
| \(t\) | 控制周期索引 | \(t=0,1,\ldots,T\)，周期可为 0.02 s |
| \(x_t\) | 第 \(t\) 个周期的车辆状态 | 包括位置、航向、速度、横摆角速度等 |
| \(x^{\mathrm{mech}}_{t+1}\) | 机理动力学模型预测的下一周期状态 | 与 \(x_t\) 同维度 |
| \(r_t\) | 第 \(t\) 个周期查询到的参考轨迹状态 | 包括参考位置、航向、曲率和速度 |
| \(u_t\) | 第 \(t\) 个周期的控制指令 | 包括转向、加速度或扭矩 |
| \(\theta\) | 待整定控制器参数集合 | 含横向参数 \(\theta_{\mathrm{lat}}\) 与纵向参数 \(\theta_{\mathrm{lon}}\) |
| \(\phi\) | 车辆机理动力学参数集合 | 质量、侧偏刚度、轮胎半径等 |
| \(h_{\psi}(\cdot)\) | MLP 残差模型 | \(\psi\) 为冻结的网络权重 |
| \(z_t\) | MLP 输入特征 | 由车辆状态、控制指令、车辆配置和归一化统计量构成 |
| \(\mathcal{T}(\cdot)\) | 残差转换算子 | 将 MLP 输出转换为车辆状态修正 |
| \(d_t^{\mathrm{ood}}\) | MLP 输入分布外距离 | 可由归一化特征的最大偏离或均方距离得到 |
| \(e_{t,\mathrm{lat}}\) | 横向跟踪误差 | 单位 m |
| \(e_{t,\mathrm{head}}\) | 航向误差 | 单位 rad |
| \(e_{t,\mathrm{spd}}\) | 速度误差 | 单位 m/s |
| \(\Delta u_t\) | 相邻周期控制指令变化 | 用于平滑性约束 |
| \(J(\theta)\) | 综合损失函数 | 无量纲或按归一化后加权 |
| \(\Pi_{\Theta}(\cdot)\) | 参数投影算子 | 将参数限制在集合 \(\Theta\) 内 |

#### （2）闭环状态更新

控制器和混合被控对象在每一周期形成如下关系。式（1）为控制器映射：

\[
u_t = g(x_t,r_t,\theta) \qquad \mathrm{(1)}
\]

式（2）为机理动力学模型预测：

\[
x^{\mathrm{mech}}_{t+1}=f_{\mathrm{mech}}(x_t,u_t,\phi) \qquad \mathrm{(2)}
\]

式（3）为 MLP 残差修正后的混合被控对象状态更新：

\[
x_{t+1}=x^{\mathrm{mech}}_{t+1}+\mathcal{T}\left(h_{\psi}(z_t)\right) \qquad \mathrm{(3)}
\]

其中，\(g(\cdot)\) 表示双模式控制器在训练路径中的可微形式，\(f_{\mathrm{mech}}(\cdot)\) 表示车辆机理动力学模型，\(h_{\psi}(\cdot)\) 表示冻结权重的 MLP 残差模型。训练时，MLP 残差模型参与车辆状态计算，但其权重不作为控制器整定变量；车辆状态误差仍可沿时间链路对 \(\theta\) 求梯度。

#### （3）综合损失

训练阶段可采用如下综合损失。式（4）为多目标加权损失：

\[
\begin{aligned}
J(\theta)=&\sum_{t=0}^T\left(w_{\mathrm{lat}}e_{t,\mathrm{lat}}^2+w_{\mathrm{head}}e_{t,\mathrm{head}}^2+w_{\mathrm{spd}}e_{t,\mathrm{spd}}^2+w_{\mathrm{smooth}}\lVert\Delta u_t\rVert^2\right)\\
&+w_{\mathrm{reg}}\lVert\theta-\theta_0\rVert^2 \qquad \mathrm{(4)}
\end{aligned}
\]

其中 \(\theta_0\) 为初始工程标定参数。上述损失可以按轨迹长度、速度段和车辆参数域进行归一化，避免长轨迹或高误差场景支配训练。

#### （4）参数更新、硬逻辑复验与诊断

参数更新可表示为式（5）：

\[
\theta^{k+1}=\Pi_{\Theta}\left(\theta^k-\eta\nabla_{\theta}J(\theta^k)\right) \qquad \mathrm{(5)}
\]

训练结束得到 \(\theta^*\) 后，将其放入验证模式，如式（6）：

\[
M_{\mathrm{hard}}(\theta^*) \rightarrow \{\mathrm{trajectory},\mathrm{error},\mathrm{command},\mathrm{log},\mathrm{diagnosis}\} \qquad \mathrm{(6)}
\]

式（6）表示用原始硬分支、硬限幅和硬速率限制运行全场景验证，并输出轨迹、误差、控制指令、日志和必要的 MLP 诊断材料。只有当验证结果满足预设要求，且诊断结果未显示 MLP 残差模型存在持续偏置、分布外输入或残差饱和等异常时，整定参数才作为可部署配置输出。

"""
    md = replace_section(md, "### 3.4.1 符号与公式", "### 3.5 关键技术参数", section341)

    section35 = """### 3.5 关键技术参数

| 符号 | 参数内容 | 约束或范围 | 是否参与整定 | 作用 |
|------|----------|------------|--------------|------|
| \\(\\theta_{\\mathrm{lat}}\\) | 横向预瞄时间、收敛时间、角速度误差预瞄时间、远预瞄时间查找表节点 | 受速度段和转向稳定性约束 | 可调 | 对应横向控制器跟踪精度和转向平滑性 |
| \\(\\theta_{\\mathrm{lon}}\\) | 站位环增益、低速速度环增益、高速速度环增益、低高速切换速度 | 增益非负，切换速度位于工程范围内 | 可调 | 对应速度跟踪和加减速响应 |
| \\(\\phi\\) | 质量、前后轴侧偏刚度、轮胎半径、传动效率、牵引车-挂车铰接参数 | 名义值或采样范围 | 固定或采样 | 用于机理动力学模型与域随机化 |
| \\(\\psi\\) | MLP 残差模型权重 | 由离线训练或高保真数据拟合得到 | 固定 | 补偿机理模型与真实车辆响应之间的剩余误差 |
| \\(z_t\\) | MLP 输入特征 | 与训练时特征定义和归一化统计一致 | 固定规则 | 决定 MLP 残差模型的输入分布 |
| \\(\\mathcal{T}\\) | 残差转换与限幅规则 | 包括坐标转换、输出裁剪和无挂车掩码等 | 固定 | 将网络输出转化为车辆状态修正 |
| \\(d_t^{\\mathrm{ood}}\\) | MLP 输入分布外距离阈值 | 可按归一化特征阈值或统计分位数确定 | 预设 | 用于诊断 MLP 是否进入训练分布外区域 |
| \\(w_{\\mathrm{lat}},w_{\\mathrm{head}},w_{\\mathrm{spd}}\\) | 跟踪误差权重 | 正数 | 预设 | 用于平衡横向、航向和速度目标 |
| \\(w_{\\mathrm{smooth}}\\) | 控制指令平滑权重 | 正数 | 预设 | 用于抑制转向和加速度突变 |
| \\(\\Pi_{\\Theta}\\) | 参数投影范围 | 上下界由车辆物理和安全约束确定 | 固定 | 防止自动更新越过可部署范围 |

"""
    md = replace_section(md, "### 3.5 关键技术参数", "## 四、与现有技术相比的有益效果", section35)

    section4 = """## 四、与现有技术相比的有益效果

本发明相较于现有技术至少具有以下有益效果。

（1）本发明不是单纯在线查表调节横向参数，也不是单独优化纵向油门或制动控制量，而是将横向和纵向工程控制器放入同一个车辆动力学闭环中进行联合整定。

（2）本发明通过双模式控制器解决可微训练与工程验证之间的矛盾。训练模式可对非光滑环节做平滑近似以获得梯度，验证模式仍保留原始硬限幅、硬分支和硬速率限制。

（3）本发明能够把多轨迹、多速度和车辆物理参数不确定性纳入同一批量训练过程，降低参数只适配单一车辆或单一场景的风险。

（4）本发明通过机理动力学模型与 MLP 残差模型组合的混合被控对象，提高闭环仿真对真实车辆响应或高保真仿真的贴近程度，同时保留机理模型的可解释性和物理边界。

（5）本发明在控制器整定阶段冻结 MLP 残差模型权重，只把 MLP 作为被控对象状态更新的一部分使用，避免将车辆模型训练与控制器参数整定混为一个不可追溯的联合优化问题。

（6）本发明通过综合损失同时约束横向误差、航向误差、速度误差、控制平滑性和参数偏移，避免只追求单一误差而牺牲控制稳定性。

（7）本发明通过参数投影模块把自动更新后的参数限制在工程可部署范围内，使输出参数更容易进入原控制模块配置文件。

（8）本发明能够输出训练曲线、参数变化图、硬逻辑验证统计、代表场景对比图和 MLP 可视化诊断图，既便于追溯参数整定过程，也便于判断异常来自控制器还是来自车辆模型残差。

"""
    md = replace_section(md, "## 四、与现有技术相比的有益效果", "## 五、本发明的技术关键点和保护点", section4)

    section5 = """## 五、本发明的技术关键点和保护点

本发明建议重点保护以下技术方案和技术特征。

（1）保护一种基于车辆动力学可微闭环仿真的车辆横纵向控制器参数自动整定方法，其至少包括：读取工程控制器逻辑和参数配置、构建双模式控制器、构建混合被控对象、批量闭环展开多场景轨迹、计算综合损失、反向传播更新参数、执行物理约束投影、硬逻辑复验并导出整定参数。

（2）保护双模式控制器结构。该结构使训练路径采用可微近似，验证路径保留原始硬逻辑，两条路径共享同一组待整定参数。

（3）保护由机理动力学模型与 MLP 残差模型组成的混合被控对象。机理模型给出车辆状态的主体预测，MLP 残差模型输出速度残差、位姿残差或二者组合，经坐标转换和限幅后叠加到机理预测状态上。

（4）保护 MLP 残差模型冻结而控制器参数可调的整定方式。训练阶段允许梯度穿过 MLP 残差模型输入输出链路回到控制器参数，但不更新 MLP 模型权重，从而把车辆响应补偿和控制器参数整定分离。

（5）保护横向控制器和纵向控制器串行闭环展开的整定链路。横向控制结果、纵向控制结果、机理动力学预测和 MLP 残差修正按控制周期串联，使多个周期后的轨迹误差能够反向影响控制器参数。

（6）保护参数集合的自动识别与分组整定方式。待整定参数可包括横向预瞄类参数、收敛类参数、角速度误差预瞄类参数、远预瞄查表参数、纵向位置环参数、纵向速度环参数和速度段切换参数。

（7）保护面向车辆物理参数不确定性的域随机化整定方式。训练时可对车辆质量、侧偏刚度、轮胎或传动相关参数采样，使整定结果兼顾多组车辆物理参数。

（8）保护反馈噪声和指令抖动参与训练的鲁棒整定方式。训练时可对位置、航向、速度、横摆角速度等反馈量加入噪声，并对转向、扭矩或加速度指令加入扰动。

（9）保护综合损失构造方式。综合损失至少包括横向跟踪误差、航向误差、速度误差、控制指令平滑项和参数偏移正则项中的一种或多种组合。

（10）保护参数投影和可部署配置生成方式。参数更新后通过投影算子限制在预设范围内，并导出与原控制器配置格式对应的参数文件。

（11）保护硬逻辑复验方式。整定完成后将参数放回保留原始硬限幅、硬分支和硬速率限制的验证路径中运行全场景闭环仿真，以验证部署前性能。

（12）保护 MLP 可视化诊断方式。该方式捕获 MLP 输入、归一化距离、残差输出、车辆状态和控制指令，生成开环静态扫描、闭环时序、输入分布外距离、危险区热图、组件消融和跨场景汇总图，用于判断闭环异常主要来源于控制器因素还是车辆动力学模型残差失效。

"""
    md = replace_section(md, "## 五、本发明的技术关键点和保护点", "## 六、实施例、技术效果和参数示例", section5)

    md = md.replace(
        "在一个实施例中，横向控制器采用重卡横向控制逻辑，纵向控制器包含位置环、速度环和扭矩输出层。车辆模型采用牵引车-挂车运动学或动力学近似模型。训练集包含换道、双换道、渐变曲率弯道、S弯、弯前减速和换道加减速等 48 条轨迹，并覆盖 5 km/h 至 50 km/h 的多个速度段。",
        "在一个实施例中，横向控制器采用重卡横向控制逻辑，纵向控制器包含位置环、速度环和扭矩输出层。车辆模型采用牵引车-挂车机理动力学模型，并可叠加 MLP 残差模型形成混合被控对象。训练集包含换道、双换道、渐变曲率弯道、S弯、弯前减速和换道加减速等 48 条轨迹，并覆盖 5 km/h 至 50 km/h 的多个速度段。"
    )
    section64_v6 = """### 6.4 实施例四：MLP 残差模型可视化诊断

在另一个实施例中，系统针对混合被控对象中的 MLP 残差模型运行可视化诊断。诊断时，系统在不改变控制器主体结构的情况下，分别运行纯机理模型路径、机理模型加完整 MLP 残差路径，以及屏蔽部分 MLP 输出分量的消融路径。每个路径均记录车辆轨迹、横向误差、方向盘命令、MLP 输入特征、归一化输入距离、原始残差输出和限幅后的残差输出。

在开环静态扫描中，系统向 MLP 残差模型输入无侧向速度、无横摆角速度和无控制激励的干净车辆状态，并扫描纵向速度。若 MLP 在该输入下仍输出随车速持续增长的侧向速度残差，则说明残差模型可能存在无激励偏置。该偏置在 50 Hz 闭环中会被逐步积分放大，使车辆状态偏离参考路径，并诱发控制器进行反向补偿。

在组件消融中，系统可以将 MLP 输出中的牵引车速度残差、挂车速度残差或相对位姿残差分别置零，再重新运行相同场景。若屏蔽某一类残差后，轨迹误差恢复到纯机理模型水平，则可判断失控主要由该类 MLP 残差输出驱动；若纯机理模型和混合模型均存在类似误差，则更可能是控制器参数或控制器结构不足。

图8示出了一个 MLP 残差可视化诊断示例。该示例通过开环静态扫描、闭环输出时序、车辆侧向速度变化、方向盘反应、早期积分偏差、轨迹平面和组件消融结果，把“网络残差偏置被高频闭环积分放大”的过程串联起来。该诊断结果用于辅助判断异常来源，不限定本发明必须采用图中具体场景、网络编号、阈值或显示样式。

![图8 MLP 残差可视化诊断示例图](figures/fig8_mlp_diagnostic_story.png)

### 6.5 可实施性说明

上述实施例可以采用通用计算设备、仿真服务器或车端开发环境实施。实施时，将车辆横向控制逻辑、纵向控制逻辑、车辆机理动力学模型、MLP 残差模型、训练场景库、参数边界表、诊断阈值和验证场景库配置在同一整定系统中；训练完成后，系统输出与原控制器参数格式一致或可转换为原参数格式的整定配置。

在一个可选实施方式中，系统保存每轮训练的参数变化、损失变化、代表场景跟踪曲线、硬逻辑复验统计、MLP 可视化诊断图、输入分布外统计和配置导出日志。上述材料用于说明参数整定过程、复验结果、车辆模型可信度和部署前差异，不限定本发明的具体软件目录、编程语言、配置文件名称或图表样式。
"""
    if "### 6.4 可实施性说明" in md:
        md = md[: md.index("### 6.4 可实施性说明")] + section64_v6
    return md


def add_equation_numbers(md: str) -> str:
    """Add visible equation numbers to the displayed formulas cited in text."""
    replacements = {
        r"""
\[
u_t = g(x_t,r_t,\theta)
\]
""".strip(): r"""
\[
u_t = g(x_t,r_t,\theta) \qquad \mathrm{(1)}
\]
""".strip(),
        r"""
\[
x_{t+1} = f(x_t,u_t,\phi)
\]
""".strip(): r"""
\[
x_{t+1} = f(x_t,u_t,\phi) \qquad \mathrm{(2)}
\]
""".strip(),
        r"""
\[
J(\theta)=\sum_{t=0}^T\left(w_{\mathrm{lat}}e_{t,\mathrm{lat}}^2+w_{\mathrm{head}}e_{t,\mathrm{head}}^2+w_{\mathrm{spd}}e_{t,\mathrm{spd}}^2+w_{\mathrm{smooth}}\lVert\Delta u_t\rVert^2\right)+w_{\mathrm{reg}}\lVert\theta-\theta_0\rVert^2
\]
""".strip(): r"""
\[
\begin{aligned}
J(\theta)=&\sum_{t=0}^T\left(w_{\mathrm{lat}}e_{t,\mathrm{lat}}^2+w_{\mathrm{head}}e_{t,\mathrm{head}}^2+w_{\mathrm{spd}}e_{t,\mathrm{spd}}^2+w_{\mathrm{smooth}}\lVert\Delta u_t\rVert^2\right)\\
&+w_{\mathrm{reg}}\lVert\theta-\theta_0\rVert^2 \qquad \mathrm{(3)}
\end{aligned}
\]
""".strip(),
        r"""
\[
\theta^{k+1}=\Pi_{\Theta}\left(\theta^k-\eta\nabla_{\theta}J(\theta^k)\right)
\]
""".strip(): r"""
\[
\theta^{k+1}=\Pi_{\Theta}\left(\theta^k-\eta\nabla_{\theta}J(\theta^k)\right) \qquad \mathrm{(4)}
\]
""".strip(),
        r"""
\[
M_{\mathrm{hard}}(\theta^*) \rightarrow \{\mathrm{trajectory},\mathrm{error},\mathrm{command},\mathrm{log}\}
\]
""".strip(): r"""
\[
M_{\mathrm{hard}}(\theta^*) \rightarrow \{\mathrm{trajectory},\mathrm{error},\mathrm{command},\mathrm{log}\} \qquad \mathrm{(5)}
\]
""".strip(),
    }
    for old, new in replacements.items():
        md = md.replace(old, new)
    return md


def disclosure_md(timestamp: str) -> str:
    return formalize_markdown(V3.disclosure_md(timestamp))


def docx_markdown(md: str) -> str:
    md = md.replace(SYSTEM_MERMAID.strip(), "![图1 可微闭环整定系统框图](figures/fig1_system_architecture.png)")
    md = md.replace(FLOW_MERMAID.strip(), "![图2 控制器参数自动整定流程](figures/fig2_method_flow.png)")
    return md


def run_pandoc(src_md: Path, out_docx: Path) -> None:
    cmd = [
        "pandoc",
        str(src_md),
        "-o",
        str(out_docx),
        "--from",
        "markdown+tex_math_dollars+tex_math_single_backslash",
        "--resource-path",
        str(OUT_DIR),
        "--metadata",
        f"title={CASE_NAME}",
    ]
    subprocess.run(cmd, check=True, cwd=OUT_DIR)


def clear_heading_italics(path: Path) -> None:
    from docx import Document

    doc = Document(path)
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            for run in para.runs:
                run.italic = False
    doc.save(path)


def apply_three_line_tables(path: Path) -> None:
    """Convert all Word tables to a standard academic three-line style."""
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    def ensure_child(parent, tag: str):
        child = parent.find(qn(tag))
        if child is None:
            child = OxmlElement(tag)
            parent.append(child)
        return child

    def set_border(parent, edge: str, val: str, size: str = "0") -> None:
        border = ensure_child(parent, f"w:{edge}")
        border.set(qn("w:val"), val)
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")

    def set_cell_margin(cell, top: int = 80, left: int = 120, bottom: int = 80, right: int = 120) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_mar = tc_pr.find(qn("w:tcMar"))
        if tc_mar is None:
            tc_mar = OxmlElement("w:tcMar")
            tc_pr.append(tc_mar)
        for edge, value in {"top": top, "left": left, "bottom": bottom, "right": right}.items():
            margin = ensure_child(tc_mar, f"w:{edge}")
            margin.set(qn("w:w"), str(value))
            margin.set(qn("w:type"), "dxa")

    doc = Document(path)
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        tbl_pr = table._tbl.tblPr
        tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
        if tbl_borders is None:
            tbl_borders = OxmlElement("w:tblBorders")
            tbl_pr.append(tbl_borders)
        for edge in ["left", "right", "insideH", "insideV"]:
            set_border(tbl_borders, edge, "nil")
        set_border(tbl_borders, "top", "single", "12")
        set_border(tbl_borders, "bottom", "single", "12")

        row_count = len(table.rows)
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_borders = tc_pr.find(qn("w:tcBorders"))
                if tc_borders is None:
                    tc_borders = OxmlElement("w:tcBorders")
                    tc_pr.append(tc_borders)

                for edge in ["left", "right", "insideH", "insideV", "top", "bottom"]:
                    set_border(tc_borders, edge, "nil")
                if row_idx == 0:
                    set_border(tc_borders, "top", "single", "12")
                    set_border(tc_borders, "bottom", "single", "8")
                if row_idx == row_count - 1:
                    set_border(tc_borders, "bottom", "single", "12")

                set_cell_margin(cell)
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_idx == 0 or col_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                    para.paragraph_format.line_spacing = 1.2
                    para.paragraph_format.space_after = Pt(2)
                    for run in para.runs:
                        V3.set_run_font(run, "宋体", 10.5, True if row_idx == 0 else None)

    doc.save(path)


def build(timestamp: str | None = None) -> tuple[Path, Path, dict[str, int]]:
    timestamp = timestamp or datetime.now().strftime("%Y%m%d%H%M%S")
    generate_figures()
    TMP_DIR.mkdir(parents=True, exist_ok=True)

    base = f"{V3.safe_name(CASE_NAME)}_{timestamp}"
    md_path = OUT_DIR / f"{base}.md"
    docx_path = OUT_DIR / f"{base}.docx"
    tmp_md = TMP_DIR / f"{base}.docx.md"

    md = disclosure_md(timestamp)
    md_path.write_text(md, encoding="utf-8")
    tmp_md.write_text(docx_markdown(md), encoding="utf-8")

    run_pandoc(tmp_md, docx_path)
    V3.style_docx(docx_path)
    clear_heading_italics(docx_path)
    apply_three_line_tables(docx_path)
    stats = V3.patch_docx_ooxml(docx_path)
    return md_path, docx_path, stats


if __name__ == "__main__":
    md, docx, stats = build()
    print(f"wrote {md}")
    print(f"wrote {docx}")
    print(f"stats {stats}")
