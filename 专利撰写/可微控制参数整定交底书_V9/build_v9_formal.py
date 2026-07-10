from __future__ import annotations

import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt
from PIL import Image, ImageDraw, ImageFont


HERE = Path(__file__).resolve().parent
DRAFT = HERE / "draft_agent_formal.md"
TMP = HERE / "_tmp"
FIGURES = HERE / "figures"
MERMAID_DIR = FIGURES / "mermaid"
TIMESTAMP = "20260710123458"
TITLE = "一种基于车辆动力学可微闭环仿真的车辆横纵向控制器参数自动整定方法"
FINAL_MD = HERE / f"{TITLE}_{TIMESTAMP}.md"
FINAL_DOCX = HERE / f"{TITLE}_{TIMESTAMP}.docx"
NORMALIZED_MD = TMP / "normalized.md"
PANDOC_MD = TMP / "pandoc_input.md"
RAW_DOCX = TMP / "pandoc_raw.docx"
PYTHON = Path(
    r"C:\Users\WuZhengc\.cache\codex-runtimes\codex-primary-runtime"
    r"\dependencies\python\python.exe"
)
MERMAID_RENDER = Path(
    r"C:\Users\WuZhengc\.codex\skills\patent-disclosure-skill"
    r"\tools\mermaid_render.py"
)


def set_run_font(run, size: float = 10.5, *, bold=None, italic=None, east_asia="宋体"):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), east_asia)


def set_style_font(style, size: float, *, bold=False, east_asia="宋体"):
    style.font.name = "Times New Roman"
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = False
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), east_asia)


def set_cell_margins(cell, top=60, start=80, bottom=60, end=80):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, *, grid=False):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        if grid or edge in ("top", "bottom"):
            node.set(qn("w:val"), "single")
            node.set(qn("w:sz"), "8" if grid else "12")
            node.set(qn("w:color"), "000000")
            node.set(qn("w:space"), "0")
        else:
            node.set(qn("w:val"), "nil")


def set_cell_bottom_border(cell, size=8):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:color"), "000000")
    bottom.set(qn("w:space"), "0")


def clear_paragraph(paragraph):
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def set_cell_text(cell, text, *, align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5, bold=False):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    set_run_font(run, size, bold=bold)


def set_first_cell_label(cell, prefix, small, suffix):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    for text, size in ((prefix, 10.5), (small, 9.0), (suffix, 10.5)):
        set_run_font(paragraph.add_run(text), size)


def make_summary_figure():
    FIGURES.mkdir(parents=True, exist_ok=True)
    font_path = Path(r"C:\Windows\Fonts\msyh.ttc")
    if not font_path.exists():
        raise FileNotFoundError(font_path)
    fonts = {
        "title": ImageFont.truetype(str(font_path), 38),
        "axis": ImageFont.truetype(str(font_path), 25),
        "tick": ImageFont.truetype(str(font_path), 22),
        "small": ImageFont.truetype(str(font_path), 20),
    }
    image = Image.new("RGB", (1980, 760), "white")
    draw = ImageDraw.Draw(image)
    colors = ("#6F8FAF", "#C89F72")
    labels = ["纯机理\n多车辆域", "冻结残差\n模型", "参数域+噪声\n与抖动"]

    def centered_text(x, y, text, font, fill="#222222"):
        box = draw.multiline_textbbox((0, 0), text, font=font, align="center", spacing=2)
        width = box[2] - box[0]
        draw.multiline_text((x - width / 2, y), text, font=font, fill=fill, align="center", spacing=2)

    def panel(origin_x, title, ylabel, maximum, series, legends, tick_step):
        left = origin_x + 105
        top = 105
        right = origin_x + 900
        bottom = 610
        centered_text((left + right) / 2, 25, title, fonts["title"])
        draw.line((left, top, left, bottom), fill="#333333", width=3)
        draw.line((left, bottom, right, bottom), fill="#333333", width=3)
        y = 0.0
        while y <= maximum + 1e-9:
            py = bottom - (y / maximum) * (bottom - top)
            draw.line((left, py, right, py), fill="#D8DDE3", width=1)
            label = f"{int(y)}" if float(y).is_integer() else f"{y:.1f}"
            box = draw.textbbox((0, 0), label, font=fonts["tick"])
            draw.text((left - 14 - (box[2] - box[0]), py - 13), label, font=fonts["tick"], fill="#444444")
            y += tick_step

        bar_width = 58
        group_centers = [left + 150, left + 395, left + 640]
        for group_index, center in enumerate(group_centers):
            for series_index, values in enumerate(series):
                value = values[group_index]
                x0 = center + (series_index - 0.5) * (bar_width + 8) - bar_width / 2
                x1 = x0 + bar_width
                y0 = bottom - (value / maximum) * (bottom - top)
                draw.rectangle((x0, y0, x1, bottom), fill=colors[series_index])
                value_text = f"{value:.2f}" if maximum < 10 else f"{int(value)}"
                centered_text((x0 + x1) / 2, y0 - 30, value_text, fonts["small"])
            centered_text(center, bottom + 16, labels[group_index], fonts["small"])

        legend_x = right - 255
        legend_y = top + 8
        for index, legend in enumerate(legends):
            y0 = legend_y + index * 38
            draw.rectangle((legend_x, y0, legend_x + 28, y0 + 20), fill=colors[index])
            draw.text((legend_x + 38, y0 - 5), legend, font=fonts["small"], fill="#333333")

        vertical = Image.new("RGBA", (430, 48), (255, 255, 255, 0))
        vdraw = ImageDraw.Draw(vertical)
        box = vdraw.textbbox((0, 0), ylabel, font=fonts["axis"])
        vdraw.text(((430 - (box[2] - box[0])) / 2, 6), ylabel, font=fonts["axis"], fill="#333333")
        vertical = vertical.rotate(90, expand=True)
        image.paste(vertical, (origin_x + 12, 145), vertical)

    panel(
        0,
        "训练损失变化",
        "归一化综合损失",
        5.0,
        ([2.7083, 3.9303, 4.6378], [1.5639, 3.0535, 3.5931]),
        ("训练前", "训练后"),
        1.0,
    )
    panel(
        990,
        "硬逻辑复验改善场景数",
        "改善场景数（总计49）",
        50.0,
        ([47, 38, 43], [47, 43, 33]),
        ("横向误差改善", "航向误差改善"),
        10.0,
    )
    image.save(FIGURES / "fig3_experiment_summary.png", dpi=(220, 220))


def make_diagram_figures():
    MERMAID_DIR.mkdir(parents=True, exist_ok=True)
    font_path = Path(r"C:\Windows\Fonts\msyh.ttc")
    title_font = ImageFont.truetype(str(font_path), 36)
    body_font = ImageFont.truetype(str(font_path), 31)
    label_font = ImageFont.truetype(str(font_path), 25)

    def centered(draw, box, text, font, fill="#1F2933"):
        x0, y0, x1, y1 = box
        bounds = draw.multiline_textbbox((0, 0), text, font=font, align="center", spacing=5)
        width = bounds[2] - bounds[0]
        height = bounds[3] - bounds[1]
        draw.multiline_text(
            ((x0 + x1 - width) / 2, (y0 + y1 - height) / 2 - 4),
            text,
            font=font,
            fill=fill,
            align="center",
            spacing=5,
        )

    def box(draw, rect, text, fill="#EEF4F8", outline="#365A73", font=None):
        draw.rounded_rectangle(rect, radius=18, fill=fill, outline=outline, width=4)
        centered(draw, rect, text, font or body_font)

    def arrow(draw, start, end, color="#365A73", width=5):
        x0, y0 = start
        x1, y1 = end
        draw.line((x0, y0, x1, y1), fill=color, width=width)
        dx, dy = x1 - x0, y1 - y0
        length = max((dx * dx + dy * dy) ** 0.5, 1)
        ux, uy = dx / length, dy / length
        px, py = -uy, ux
        size = 18
        base_x, base_y = x1 - ux * size, y1 - uy * size
        points = [
            (x1, y1),
            (base_x + px * size * 0.55, base_y + py * size * 0.55),
            (base_x - px * size * 0.55, base_y - py * size * 0.55),
        ]
        draw.polygon(points, fill=color)

    system = Image.new("RGB", (2200, 1180), "white")
    draw = ImageDraw.Draw(system)
    boxes = {
        "A": (70, 70, 410, 220),
        "B": (500, 70, 840, 220),
        "C": (930, 70, 1270, 220),
        "D": (70, 370, 410, 530),
        "F": (500, 370, 840, 530),
        "E": (930, 370, 1270, 530),
        "H": (1360, 370, 1700, 530),
        "I": (1790, 370, 2130, 530),
        "G": (500, 720, 840, 880),
        "L": (930, 720, 1270, 880),
        "K": (1360, 720, 1700, 880),
        "J": (1790, 720, 2130, 880),
        "M": (1360, 990, 1700, 1130),
    }
    texts = {
        "A": "工程控制器逻辑\n和参数配置",
        "B": "参数登记\n及边界确定",
        "C": "共享参数的\n双路径控制器",
        "D": "多轨迹、多速度\n和车辆参数域",
        "F": "机理车辆\n动力学模型",
        "E": "可微闭环\n时间展开",
        "H": "逐场景归一化\n损失",
        "I": "自动微分\n和参数更新",
        "G": "可选：冻结的\n残差模型",
        "J": "参数投影",
        "K": "原始硬逻辑\n闭环复验",
        "L": "配置与报告\n导出",
        "M": "可选：残差模型\n可信度诊断",
    }
    for key, rect in boxes.items():
        optional = key in {"G", "M"}
        box(
            draw,
            rect,
            texts[key],
            fill="#F5F5F5" if optional else "#EEF4F8",
            outline="#777777" if optional else "#365A73",
        )
    arrow(draw, (410, 145), (500, 145))
    arrow(draw, (840, 145), (930, 145))
    arrow(draw, (1100, 220), (1100, 370))
    arrow(draw, (410, 450), (930, 450))
    arrow(draw, (840, 450), (930, 450))
    arrow(draw, (670, 720), (670, 530), color="#777777")
    arrow(draw, (1270, 450), (1360, 450))
    arrow(draw, (1700, 450), (1790, 450))
    arrow(draw, (1960, 530), (1960, 720))
    arrow(draw, (1790, 800), (1700, 800))
    arrow(draw, (1360, 800), (1270, 800))
    draw.line((1440, 720, 1440, 620, 1100, 620), fill="#A05A55", width=5)
    arrow(draw, (1100, 620), (1100, 530), color="#A05A55")
    draw.text((1180, 582), "未达标", font=label_font, fill="#A05A55")
    draw.text((1278, 756), "达标", font=label_font, fill="#365A73")
    arrow(draw, (1530, 880), (1530, 990), color="#777777")
    draw.text((1550, 920), "模型异常", font=label_font, fill="#666666")
    system.save(MERMAID_DIR / "fig1_system.png", dpi=(220, 220))

    flow = Image.new("RGB", (1800, 1540), "white")
    draw = ImageDraw.Draw(flow)
    step_boxes = [
        (380, 40, 1420, 150),
        (380, 215, 1420, 325),
        (380, 390, 1420, 500),
        (380, 565, 1420, 675),
        (380, 740, 1420, 850),
        (380, 915, 1420, 1025),
        (380, 1090, 1420, 1200),
        (380, 1370, 1420, 1480),
    ]
    step_texts = [
        "S1  取得工程控制器参数、边界和回写映射",
        "S2  构建共享参数的训练路径与验证路径",
        "S3  配置车辆动力学、轨迹批次和车辆参数域",
        "S4  按横向—纵向—车辆状态顺序闭环展开",
        "S5  计算并归一化各场景综合损失",
        "S6  反向传播、更新参数并执行边界投影",
        "S7  使用原始硬逻辑运行独立闭环复验",
        "S8  回写参数配置并输出验证报告",
    ]
    for index, (rect, text) in enumerate(zip(step_boxes, step_texts)):
        fill = "#EDF6F0" if index == 7 else "#EEF4F8"
        box(draw, rect, text, fill=fill, font=body_font)
    for index in range(6):
        arrow(draw, (900, step_boxes[index][3]), (900, step_boxes[index + 1][1]))
    arrow(draw, (900, step_boxes[6][3]), (900, step_boxes[7][1]))
    draw.text((925, 1270), "通过", font=label_font, fill="#365A73")

    diagnostic = (1460, 1090, 1760, 1260)
    box(
        draw,
        diagnostic,
        "可选诊断\n静态扫描、分布距离\n和组件消融",
        fill="#F5F5F5",
        outline="#777777",
        font=label_font,
    )
    arrow(draw, (1420, 1145), (1460, 1145), color="#777777")
    draw.text((1450, 1095), "异常", font=label_font, fill="#666666")

    draw.line((380, 1145, 170, 1145, 170, 620, 380, 620), fill="#A05A55", width=5)
    arrow(draw, (170, 620), (380, 620), color="#A05A55")
    draw.text((190, 1080), "未通过，返回闭环展开", font=label_font, fill="#A05A55")
    flow.save(MERMAID_DIR / "fig2_flow.png", dpi=(220, 220))


def normalize_markdown():
    text = DRAFT.read_text(encoding="utf-8")
    replacements = {
        "(v_{b,t})": "$v_{b,t}$",
        "(v_{\\mathrm{sw}})": "$v_{\\mathrm{sw}}$",
        "(v_{mathrm{sw}})": "$v_{\\mathrm{sw}}$",
        "(\tau>0)": "$\\tau>0$",
        "(y_{\\mathrm{low}})": "$y_{\\mathrm{low}}$",
        "(y_{\\mathrm{high}})": "$y_{\\mathrm{high}}$",
        "(y_{mathrm{low}})": "$y_{\\mathrm{low}}$",
        "(y_{mathrm{high}})": "$y_{\\mathrm{high}}$",
        "(\tau)": "$\\tau$",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    NORMALIZED_MD.write_text(text, encoding="utf-8")


def render_mermaid():
    text = NORMALIZED_MD.read_text(encoding="utf-8")
    fence = chr(96) * 3
    pattern = re.escape(fence) + r"mermaid\s*\n.*?" + re.escape(fence)
    images = [
        ("figures/mermaid/fig1_system.png", "图1 可微闭环参数自动整定系统框图"),
        ("figures/mermaid/fig2_flow.png", "图2 控制器参数自动整定方法流程图"),
    ]
    index = 0

    def add_image_comment(match):
        nonlocal index
        if index >= len(images):
            raise RuntimeError("Mermaid 图数量超过预期")
        path, alt = images[index]
        index += 1
        return f"{match.group(0)}\n\n<!-- ![{alt}]({path}) -->"

    text = re.sub(pattern, add_image_comment, text, flags=re.S)
    if index != 2:
        raise RuntimeError(f"预期写入 2 个 Mermaid 图片引用，实际写入 {index} 个")
    FINAL_MD.write_text(text, encoding="utf-8")


def make_pandoc_markdown():
    text = FINAL_MD.read_text(encoding="utf-8")
    fence = chr(96) * 3
    pattern = (
        re.escape(fence)
        + r"mermaid\s*\n.*?"
        + re.escape(fence)
        + r"\s*\n\s*<!--\s*!\[[^\]]*\]\(([^)]+)\)\s*-->"
    )

    def replace_diagram(match):
        path = match.group(1).replace("\\", "/")
        return f"![]({path}){{width=95%}}"

    text, count = re.subn(pattern, replace_diagram, text, flags=re.S)
    if count != 2:
        raise RuntimeError(f"预期替换 2 个 Mermaid 图，实际替换 {count} 个")
    PANDOC_MD.write_text(text, encoding="utf-8")


def add_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def rebuild_uniform_cover_table(document):
    old_table = document.tables[0]
    table = document.add_table(rows=7, cols=5)
    old_table._tbl.addprevious(table._tbl)
    old_table._tbl.getparent().remove(old_table._tbl)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [1755, 794, 1665, 2000, 2475]
    grid_cols = table._tbl.tblGrid.gridCol_lst
    for grid_col, width in zip(grid_cols, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = width * 635
            set_cell_margins(cell)

    row0 = table.rows[0].cells
    c00 = row0[0].merge(row0[1])
    c02 = row0[2]
    c03 = row0[3].merge(row0[4])
    set_first_cell_label(c00, "第一发明人（", "必填", "）")
    set_cell_text(c02, "")
    set_cell_text(c03, "☐  校招  /  ☐  社招  （必选）")

    row1 = table.rows[1].cells
    c10 = row1[0].merge(row1[1])
    c12 = row1[2].merge(row1[4])
    c10.text = ""
    c10.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = c10.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    for text, size in (("其他发明人（", 10.5), ("不超过3人", 9.0), ("）", 10.5)):
        set_run_font(p.add_run(text), size)
    set_cell_text(c12, "")

    row2 = table.rows[2].cells
    c20 = row2[0].merge(row2[2])
    c23 = row2[3].merge(row2[4])
    set_cell_text(c20, "（以下由交底书撰写人填写）", size=9.0)
    set_cell_text(c23, "（以下由知识产权部填写）", size=9.0)

    labels = [
        ("撰写人", "专利类型"),
        ("手机", "知识产权负责人"),
        ("座机", "联系电话"),
        ("E－mail", "E－mail"),
    ]
    for row_index, (left_label, right_label) in enumerate(labels, start=3):
        cells = table.rows[row_index].cells
        middle = cells[1].merge(cells[2])
        set_cell_text(cells[0], left_label)
        set_cell_text(middle, "")
        set_cell_text(cells[3], right_label)
        set_cell_text(cells[4], "")

    heights = [395, 458, 307, 455, 455, 462, 455]
    for row, height in zip(table.rows, heights):
        row.height = height * 635
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        prevent_row_split(row)
    set_table_borders(table, grid=True)

    border_paragraph = document.add_paragraph(" ")
    table._tbl.addprevious(border_paragraph._p)
    border_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    border_paragraph.paragraph_format.line_spacing = Pt(13)
    border_paragraph.paragraph_format.space_after = Pt(18)
    p_pr = border_paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    return table


def style_hyperlinks(paragraph):
    for hyperlink in paragraph._p.xpath(".//w:hyperlink"):
        for r in hyperlink.xpath(".//w:r"):
            r_pr = r.find(qn("w:rPr"))
            if r_pr is None:
                r_pr = OxmlElement("w:rPr")
                r.insert(0, r_pr)
            color = r_pr.find(qn("w:color"))
            if color is None:
                color = OxmlElement("w:color")
                r_pr.append(color)
            color.set(qn("w:val"), "0563C1")
            underline = r_pr.find(qn("w:u"))
            if underline is None:
                underline = OxmlElement("w:u")
                r_pr.append(underline)
            underline.set(qn("w:val"), "single")
            fonts = r_pr.find(qn("w:rFonts"))
            if fonts is None:
                fonts = OxmlElement("w:rFonts")
                r_pr.insert(0, fonts)
            fonts.set(qn("w:ascii"), "Times New Roman")
            fonts.set(qn("w:hAnsi"), "Times New Roman")
            fonts.set(qn("w:eastAsia"), "宋体")
            size = r_pr.find(qn("w:sz"))
            if size is None:
                size = OxmlElement("w:sz")
                r_pr.append(size)
            size.set(qn("w:val"), "21")


def apply_document_styles(document):
    normal = document.styles["Normal"]
    set_style_font(normal, 10.5)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.first_line_indent = Pt(21)

    if "Body Text" in document.styles:
        set_style_font(document.styles["Body Text"], 10.5)
    if "List Paragraph" in document.styles:
        set_style_font(document.styles["List Paragraph"], 10.5)

    heading_specs = {
        "Heading 1": (18.0, 0, 0),
        "Heading 2": (14.0, 14, 8),
        "Heading 3": (12.0, 10, 6),
        "Heading 4": (11.0, 8, 4),
        "Heading 5": (10.5, 6, 3),
    }
    for name, (size, before, after) in heading_specs.items():
        if name not in document.styles:
            continue
        style = document.styles[name]
        set_style_font(style, size, bold=True, east_asia="黑体")
        style.font.color.rgb = None
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.12
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.keep_with_next = True

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        has_numbering = bool(paragraph._p.xpath("./w:pPr/w:numPr"))
        is_equation = bool(paragraph._p.xpath(".//m:oMathPara"))
        has_image = bool(paragraph._p.xpath(".//w:drawing"))

        for run in paragraph.runs:
            set_run_font(run)

        if text == "专利技术交底书":
            paragraph.style = normal
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            paragraph.paragraph_format.line_spacing = Pt(23)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                set_run_font(run, 18.0, bold=True)
            continue

        if text == TITLE:
            paragraph.style = normal
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(14)
            paragraph.paragraph_format.space_after = Pt(10)
            paragraph.paragraph_format.line_spacing = 1.05
            paragraph.paragraph_format.keep_with_next = True
            for run in paragraph.runs:
                set_run_font(run, 14.0, bold=True)
            continue

        if paragraph.style.name.startswith("Heading"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Pt(0)
            for run in paragraph.runs:
                level_match = re.search(r"(\d+)$", paragraph.style.name)
                level = int(level_match.group(1)) if level_match else 5
                size = {2: 14.0, 3: 12.0, 4: 11.0, 5: 10.5}.get(level, 10.5)
                set_run_font(run, size, bold=True, italic=False, east_asia="黑体")
            continue

        if is_equation:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(4)
            continue

        if has_image:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(3)
            continue

        if re.match(r"^图[123]\s", text):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(8)
            paragraph.paragraph_format.keep_with_next = False
            for run in paragraph.runs:
                set_run_font(run, 10.0, italic=False)
            continue

        if has_numbering:
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.paragraph_format.line_spacing = 1.15
        elif text:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Pt(21)
            paragraph.paragraph_format.line_spacing = 1.15
            paragraph.paragraph_format.space_after = Pt(6)

        if text.startswith("公开来源："):
            style_hyperlinks(paragraph)


def style_technical_tables(document):
    for table_index, table in enumerate(document.tables):
        if table_index == 0:
            continue
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        set_table_borders(table, grid=False)
        if table.rows:
            add_repeat_table_header(table.rows[0])
        for row_index, row in enumerate(table.rows):
            prevent_row_split(row)
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell, top=55, start=70, bottom=55, end=70)
                if row_index == 0:
                    set_cell_bottom_border(cell, size=8)
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.first_line_indent = Pt(0)
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(3)
                    paragraph.paragraph_format.line_spacing = 1.1
                    paragraph.alignment = (
                        WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
                    )
                    for run in paragraph.runs:
                        set_run_font(run, 9.5, bold=True if row_index == 0 else None)


def resize_images(document):
    max_width = Inches(6.15)
    max_heights = [Inches(3.35), Inches(3.45), Inches(2.65)]
    for index, shape in enumerate(document.inline_shapes):
        max_height = max_heights[index] if index < len(max_heights) else Inches(3.5)
        ratio = min(max_width / shape.width, max_height / shape.height, 1.0)
        shape.width = int(shape.width * ratio)
        shape.height = int(shape.height * ratio)


def set_footer_page_number(document):
    for section in document.sections:
        section.footer.is_linked_to_previous = False
        paragraph = section.footer.paragraphs[0]
        clear_paragraph(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run("第 ")
        set_run_font(run, 9.0)
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        fld_separate = OxmlElement("w:fldChar")
        fld_separate.set(qn("w:fldCharType"), "separate")
        text_run = OxmlElement("w:r")
        text_node = OxmlElement("w:t")
        text_node.text = "1"
        text_run.append(text_node)
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        for field_part in (fld_begin, instr, fld_separate):
            field_run = OxmlElement("w:r")
            field_run.append(field_part)
            paragraph._p.append(field_run)
        paragraph._p.append(text_run)
        end_run = OxmlElement("w:r")
        end_run.append(fld_end)
        paragraph._p.append(end_run)
        run = paragraph.add_run(" 页")
        set_run_font(run, 9.0)


def style_docx():
    document = Document(RAW_DOCX)
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    rebuild_uniform_cover_table(document)
    apply_document_styles(document)
    style_technical_tables(document)
    resize_images(document)
    set_footer_page_number(document)

    document.core_properties.title = TITLE
    document.core_properties.subject = "专利技术交底书代理人正式稿"
    document.core_properties.author = ""
    document.core_properties.keywords = "可微控制；参数整定；车辆动力学；横纵向控制"

    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is not None:
        settings.remove(update_fields)
    document.save(FINAL_DOCX)


def main():
    TMP.mkdir(parents=True, exist_ok=True)
    MERMAID_DIR.mkdir(parents=True, exist_ok=True)
    make_summary_figure()
    make_diagram_figures()
    normalize_markdown()
    render_mermaid()
    make_pandoc_markdown()
    subprocess.run(
        [
            "pandoc",
            str(PANDOC_MD),
            "-o",
            str(RAW_DOCX),
            "--from",
            "markdown+tex_math_dollars+tex_math_single_backslash",
            "--resource-path",
            str(HERE),
        ],
        check=True,
    )
    style_docx()
    print(FINAL_MD)
    print(FINAL_DOCX)


if __name__ == "__main__":
    main()
