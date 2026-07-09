# -*- coding: utf-8 -*-
"""Build V1 patent technical disclosure for differentiable control tuning."""

from __future__ import annotations

import shutil
from pathlib import Path
from textwrap import wrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = Path(__file__).resolve().parent
FIG_DIR = OUT_DIR / "figures"
MD_PATH = OUT_DIR / "technical_disclosure_v1.md"
DOCX_PATH = OUT_DIR / "technical_disclosure_v1.docx"

PREFERRED_TITLE = "一种基于车辆动力学可微闭环仿真的车辆横纵向控制器参数自动整定方法"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path("C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf" if bold else "C:/Windows/Fonts/simsun.ttc"),
        Path("C:/Windows/Fonts/arial.ttf"),
    ]
    for p in candidates:
        if p.exists():
            return ImageFont.truetype(str(p), size=size)
    return ImageFont.load_default()


def wrap_cjk(text: str, width: int) -> list[str]:
    lines: list[str] = []
    for part in text.split("\n"):
        if not part:
            lines.append("")
            continue
        line = ""
        for ch in part:
            if len(line) >= width:
                lines.append(line)
                line = ch
            else:
                line += ch
        if line:
            lines.append(line)
    return lines


def draw_box(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int, int, int],
    text: str,
    fill: str = "#FFFFFF",
    outline: str = "#111111",
    text_fill: str = "#111111",
    size: int = 26,
    bold: bool = False,
    wrap_width: int = 12,
) -> None:
    x0, y0, x1, y1 = xy
    draw.rounded_rectangle(xy, radius=16, fill=fill, outline=outline, width=3)
    f = font(size, bold=bold)
    lines = wrap_cjk(text, wrap_width)
    line_h = size + 8
    total_h = len(lines) * line_h
    y = y0 + (y1 - y0 - total_h) // 2 + 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=f)
        tw = bbox[2] - bbox[0]
        draw.text((x0 + (x1 - x0 - tw) // 2, y), line, font=f, fill=text_fill)
        y += line_h


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int]) -> None:
    draw.line([start, end], fill="#111111", width=4)
    x0, y0 = start
    x1, y1 = end
    if abs(x1 - x0) >= abs(y1 - y0):
        s = 1 if x1 >= x0 else -1
        pts = [(x1, y1), (x1 - 18 * s, y1 - 10), (x1 - 18 * s, y1 + 10)]
    else:
        s = 1 if y1 >= y0 else -1
        pts = [(x1, y1), (x1 - 10, y1 - 18 * s), (x1 + 10, y1 - 18 * s)]
    draw.polygon(pts, fill="#111111")


def generate_flow_figures() -> None:
    FIG_DIR.mkdir(parents=True, exist_ok=True)

    # 图1：总体流程
    img = Image.new("RGB", (1800, 1050), "white")
    d = ImageDraw.Draw(img)
    title_f = font(42, bold=True)
    d.text((70, 45), "图1 车辆横纵向控制器可微闭环参数整定总体流程", font=title_f, fill="#111111")
    boxes = {
        "a": (90, 160, 390, 285, "工业控制器\n代码/参数表"),
        "b": (510, 160, 810, 285, "双模式控制器\n可微训练/硬逻辑验证"),
        "c": (930, 160, 1230, 285, "车辆动力学\n机理/残差模型"),
        "d": (1350, 160, 1650, 285, "多场景轨迹库\n类型×速度段"),
        "e": (320, 430, 620, 555, "50Hz闭环展开\n横向→纵向→车辆"),
        "f": (750, 430, 1050, 555, "跟踪与平滑\n联合损失"),
        "g": (1180, 430, 1480, 555, "时间反传\n参数梯度"),
        "h": (540, 710, 840, 835, "参数投影\n物理约束"),
        "i": (960, 710, 1260, 835, "导出整定参数\n配置文件"),
        "j": (1380, 710, 1680, 835, "硬限幅验证\n结果图/日志"),
    }
    for key, (x0, y0, x1, y1, txt) in boxes.items():
        fill = "#EAF3F8" if key in ["b", "e", "f", "g"] else "#FFFFFF"
        draw_box(d, (x0, y0, x1, y1), txt, fill=fill, size=28, bold=key in ["b", "e"])
    arrow(d, (390, 222), (510, 222))
    arrow(d, (810, 222), (930, 222))
    arrow(d, (1230, 222), (1350, 222))
    arrow(d, (660, 285), (500, 430))
    arrow(d, (1080, 285), (500, 430))
    arrow(d, (1500, 285), (500, 430))
    arrow(d, (620, 492), (750, 492))
    arrow(d, (1050, 492), (1180, 492))
    arrow(d, (1330, 555), (700, 710))
    arrow(d, (840, 772), (960, 772))
    arrow(d, (1260, 772), (1380, 772))
    arrow(d, (1520, 710), (1510, 585))
    d.text((1320, 620), "未达标可继续整定", font=font(24), fill="#444444")
    img.save(FIG_DIR / "fig1_overall_flow.png", quality=95)

    # 图2：双路径
    img = Image.new("RGB", (1600, 900), "white")
    d = ImageDraw.Draw(img)
    d.text((70, 45), "图2 可微训练路径与硬逻辑验证路径", font=title_f, fill="#111111")
    draw_box(d, (100, 160, 440, 285), "同一控制器\n参数与状态", size=30, bold=True)
    draw_box(d, (620, 120, 1040, 260), "训练路径\n平滑分支/直通限幅\n可回传梯度", fill="#EAF3F8", size=28, bold=True, wrap_width=13)
    draw_box(d, (620, 410, 1040, 550), "验证路径\n原始硬分支/硬限幅\n对齐工程行为", fill="#F8F8F8", size=28, bold=True, wrap_width=13)
    draw_box(d, (1220, 120, 1500, 260), "参数更新\n自动整定", fill="#FFFFFF", size=28)
    draw_box(d, (1220, 410, 1500, 550), "V1对比\n轨迹/误差/指令", fill="#FFFFFF", size=28, wrap_width=9)
    arrow(d, (440, 222), (620, 190))
    arrow(d, (440, 222), (620, 480))
    arrow(d, (1040, 190), (1220, 190))
    arrow(d, (1040, 480), (1220, 480))
    arrow(d, (1360, 410), (1360, 260))
    d.text((1090, 300), "同一组整定参数\n必须在硬逻辑路径中复验",
           font=font(26, bold=True), fill="#111111", spacing=8)
    d.line((180, 680, 1420, 680), fill="#333333", width=3)
    d.text((220, 720), "非光滑处理：查找表线性插值、速率限制直通估计、条件分支平滑混合、最近点选择隔离、时间预瞄可微插值", font=font(27), fill="#222222")
    img.save(FIG_DIR / "fig2_dual_path.png", quality=95)

    # 图3：批量域随机化
    img = Image.new("RGB", (1700, 980), "white")
    d = ImageDraw.Draw(img)
    d.text((70, 45), "图3 多轨迹多域并行鲁棒整定流程", font=title_f, fill="#111111")
    draw_box(d, (80, 155, 420, 285), "48条轨迹\n8类×6速度段", size=28, bold=True, wrap_width=10)
    draw_box(d, (80, 390, 420, 520), "每轮采样K组\n车辆物理参数", size=28, bold=True, wrap_width=10)
    draw_box(d, (590, 270, 930, 400), "展开为批量\nB=48×K", fill="#EAF3F8", size=30, bold=True)
    draw_box(d, (1100, 155, 1460, 285), "同步闭环推进\n共享控制器参数", fill="#EAF3F8", size=28, bold=True, wrap_width=12)
    draw_box(d, (1100, 390, 1460, 520), "按轨迹归一化\n按域统计损失", fill="#FFFFFF", size=28, wrap_width=12)
    draw_box(d, (590, 640, 930, 770), "平均损失反传\n更新控制器", fill="#FFFFFF", size=28, wrap_width=10)
    arrow(d, (420, 220), (590, 320))
    arrow(d, (420, 455), (590, 350))
    arrow(d, (930, 335), (1100, 220))
    arrow(d, (1280, 285), (1280, 390))
    arrow(d, (1100, 455), (930, 705))
    arrow(d, (590, 705), (300, 520))
    d.text((160, 595), "下一训练轮重新采样车辆域", font=font(25), fill="#444444")
    d.text((100, 825), "可选扰动：反馈状态噪声、转向/扭矩指令抖动；验证阶段关闭扰动，走硬逻辑路径。", font=font(28), fill="#222222")
    img.save(FIG_DIR / "fig3_batched_domain_randomization.png", quality=95)


def copy_project_figures() -> dict[str, Path]:
    src = ROOT / "sim" / "results" / "training" / "truck_trailer" / "20260526_123421_mlp0525"
    mapping = {
        "fig4_loss_curve.png": src / "loss_curve.png",
        "fig5_parameter_changes.png": src / "parameter_changes.png",
        "fig6_training_summary.png": src / "training_summary.png",
        "fig7_comparison_trajectory.png": src / "comparison_trajectory.png",
        "fig8_comparison_lateral_error.png": src / "comparison_lateral_error.png",
    }
    copied: dict[str, Path] = {}
    for name, source in mapping.items():
        target = FIG_DIR / name
        if source.exists():
            if name in {"fig7_comparison_trajectory.png", "fig8_comparison_lateral_error.png"}:
                make_representative_crop(source, target)
            else:
                shutil.copy2(source, target)
            copied[name] = target
    return copied


def make_representative_crop(source: Path, target: Path, crop_height: int = 2800) -> None:
    """Use a readable representative slice of very tall all-scenario result figures."""
    with Image.open(source) as img:
        width, height = img.size
        cropped = img.crop((0, 0, width, min(crop_height, height))).convert("RGB")

    strip_h = 120
    canvas = Image.new("RGB", (cropped.width, cropped.height + strip_h), "white")
    canvas.paste(cropped, (0, strip_h))
    d = ImageDraw.Draw(canvas)
    d.text(
        (80, 34),
        "代表场景摘选；全量结果见 sim/results/training/truck_trailer/20260526_123421_mlp0525",
        font=font(42, bold=True),
        fill="#111111",
    )
    canvas.save(target, quality=95)


def disclosure_sections() -> list[dict]:
    return [
        {
            "level": 0,
            "title": "专利技术交底书 V1",
            "paras": [
                f"推荐发明名称：{PREFERRED_TITLE}",
                "版本：V1；形成日期：2026-07-09；适用材料：当前 differentiable_control 项目的代码、设计文档、训练日志和结果图。",
                "本稿定位为给专利代理人继续撰写申请文件的技术交底材料，不直接替代正式权利要求书。",
            ],
        },
        {
            "level": 1,
            "title": "0. 本稿结论与命名建议",
            "paras": [
                "本项目的核心创新不是单纯搭建一个仿真环境，而是把真实车辆横纵向控制器、车辆动力学和评价指标连接成可自动求梯度的闭环时间链路，使控制器设计参数可以从多场景跟踪误差中自动整定，并且整定后仍用原始硬限幅逻辑复验。因此主案名称建议突出“车辆动力学、可微闭环仿真、横纵向控制器参数自动整定”。",
            ],
            "tables": [
                {
                    "headers": ["序号", "名称建议", "适用策略"],
                    "rows": [
                        ["首选", PREFERRED_TITLE, "保护范围聚焦，能够覆盖项目当前最有技术效果的闭环整定流程。"],
                        ["备选1", "一种自动驾驶控制器可微仿真环境的构建及参数整定方法", "范围更宽，但“环境构建”技术效果不如参数整定明确，建议作为从属表述或拆分案。"],
                        ["备选2", "一种面向车辆动力学参数不确定性的自动驾驶控制器鲁棒可微整定方法", "适合单独保护域随机化、反馈噪声和指令抖动训练。"],
                        ["备选3", "一种工业车辆控制器双模式可微复现及硬约束验证方法", "适合保护从 C++ 控制器到可微模块、再回到硬逻辑验证的复现流程。"],
                    ],
                }
            ],
        },
        {
            "level": 1,
            "title": "1. 技术领域",
            "paras": [
                "本发明涉及车辆自动驾驶控制、车辆动力学仿真、控制器参数标定和自动微分优化技术领域，尤其涉及一种基于车辆动力学可微闭环仿真的车辆横纵向控制器参数自动整定方法、系统及计算机可读存储介质。",
            ],
        },
        {
            "level": 1,
            "title": "2. 技术背景",
            "paras": [
                "自动驾驶车辆的控制模块通常以固定频率闭环运行。规划模块给出参考轨迹和参考速度，横向控制器输出转向指令，纵向控制器输出加速度、制动或驱动扭矩，车辆再根据动力学响应进入下一时刻状态。控制器内部往往包含预瞄、查找表、级联比例积分控制、前馈补偿、限幅、速率限制、滤波和安全保护等工程逻辑。",
                "传统标定流程主要依赖工程师经验和反复试验：先在若干典型道路上调整横向预瞄时间、转向前馈、速度环增益等参数，再观察横向误差、航向误差、速度误差和指令平滑性。该流程在车辆模型变化、载荷变化、轮胎侧偏刚度变化或控制器结构变化时需要大量重复工作，且很难同时兼顾多速度段、多轨迹和多车辆状态。",
                "现有仿真工具可以批量评估控制器表现，但很多工程控制器包含不可微或近似不可微的步骤，例如最近点选择、条件分支、硬限幅、速率限制和按时间预瞄查表。常规仿真只能给出结果，不能直接告诉工程师各个控制参数应该朝哪个方向调整。黑盒优化和强化学习虽然可以搜索参数，但样本效率低、结果可解释性弱，而且难以保证调出的参数在原始硬限幅控制逻辑下仍然有效。",
            ],
        },
        {
            "level": 1,
            "title": "3. 现有方案的问题和缺陷",
            "paras": [
                "第一，人工调参难以同时覆盖横向和纵向耦合效果。横向预瞄参数会影响路径误差和转向指令，纵向速度误差又会改变车辆进入弯道时的动力学状态，两者在闭环中相互影响，单独调整某一类参数容易造成其他指标退化。",
                "第二，普通仿真评估不提供参数梯度。即使能发现某一场景横向误差变大，也不能直接得到预瞄查找表、速度环增益或切换速度等参数的调整方向。",
                "第三，工程控制器中大量硬逻辑会阻断自动微分。若直接把硬限幅、硬分支、最近点选择放入可微计算链路，会出现梯度为零、梯度爆炸或计算图断裂等问题。",
                "第四，只在单一名义车辆参数上调参容易过拟合。实车质量、侧偏刚度、挂车状态、路面附着和传感器反馈都有不确定性，名义点最优参数在偏离名义点时可能不再最优。",
                "第五，若训练时采用平滑近似而部署时采用硬限幅逻辑，二者之间可能存在行为偏差。因此需要一种整定后自动回到硬逻辑路径复验的机制，证明参数不是只适用于训练近似。"
            ],
        },
        {
            "level": 1,
            "title": "4. 发明内容和实施方式",
            "paras": [
                "本发明提供一种车辆横纵向控制器参数自动整定方法。该方法将工业车辆控制器复现为具有训练模式和验证模式的双模式控制器，将车辆动力学模型、轨迹查询、横纵向控制器和损失函数按控制周期展开为闭环时间链路，通过自动微分获得跟踪误差对控制器参数的梯度，并在物理约束下更新参数。训练结束后，将整定参数导出为可部署配置，并使用原始硬分支、硬限幅和硬速率限制的验证模式重新运行全场景仿真，输出参数变化、轨迹对比、误差对比和实验日志。",
            ],
        },
        {
            "level": 2,
            "title": "4.1 总体输入、输出和模块组成",
            "paras": [
                "本方法的输入可以包括：车辆规划轨迹、参考速度、车辆初始状态、待整定控制器参数、车辆动力学参数、训练场景集合、损失权重和参数约束。输出可以包括：整定后的控制器配置文件、训练收敛曲线、参数变化图、调参前后对比图、每个场景的误差指标和实验日志。",
                "系统可以包括以下模块：控制器双模式复现模块、轨迹生成与轨迹查询模块、车辆动力学模块、闭环仿真展开模块、损失计算模块、非光滑控制逻辑可微处理模块、参数优化与投影模块、批量鲁棒训练模块和硬逻辑验证模块。",
            ],
        },
        {
            "level": 2,
            "title": "4.2 方法步骤",
            "paras": [
                "步骤S1，获取车辆控制器的原始控制逻辑和参数配置，识别其中属于控制器设计的可调参数以及属于车辆物理、安全边界或监控输出的固定参数。",
                "步骤S2，将横向控制器和纵向控制器封装为双模式控制器。训练模式保留控制器物理含义并允许梯度通过；验证模式使用与原工程逻辑一致的硬分支、硬限幅和硬速率限制。",
                "步骤S3，构建车辆动力学被控对象。被控对象可以为运动学自行车模型、动力学自行车模型、机理模型与残差模型组合的混合模型，或牵引车与挂车双体动力学模型。对外统一输出控制器所需的车辆位置、航向、速度和横摆角速度。",
                "步骤S4，构建训练轨迹集合。训练轨迹可以覆盖换道、双换道、渐变曲率弯道、S弯、弯前减速和换道加减速等类型，并在多个速度段展开。",
                "步骤S5，在每一个控制周期内按横向控制器、纵向控制器、车辆动力学更新的顺序展开闭环仿真，形成从控制参数到车辆状态再到跟踪误差的时间计算链路。",
                "步骤S6，计算由横向误差、航向误差、速度误差、转向变化率、加速度变化率和参数偏移正则项组成的综合损失。",
                "步骤S7，沿时间链路进行反向传播，得到综合损失相对于控制器参数的梯度，利用优化器更新参数，并对预瞄时间、比例积分增益和速度切换点等参数执行物理范围投影。",
                "步骤S8，可选地，将训练轨迹复制到多个车辆参数域，在每轮训练中采样不同质量、前后轴侧偏刚度等车辆物理参数，并叠加反馈噪声或执行器指令抖动，以得到更鲁棒的参数。",
                "步骤S9，将整定参数导出为配置文件，使用验证模式在多场景上重新运行仿真，若误差、平滑性或安全约束不满足要求，则以上一轮整定结果继续训练或调整训练集合。",
            ],
        },
        {
            "level": 2,
            "title": "4.3 可调参数分类",
            "paras": [
                "在一个实施例中，横向控制器采用多点预瞄和曲率前馈结构。可调参数包括预瞄距离时间系数、横向误差收敛时间、角速度误差预瞄时间和远预瞄点时间等查找表的节点值。固定参数包括车辆几何修正、安全限幅、近预瞄监控参数、侧滑物理参数和转向速率边界。",
                "在一个实施例中，纵向控制器采用站位环和速度环级联结构。可调参数包括站位环比例和积分增益、低速速度环比例和积分增益、高速速度环比例和积分增益以及低高速切换速度。固定参数包括加速度包络、加速度变化率包络、扭矩输出层的车辆质量、风阻、滚阻、轮胎半径和传动效率等物理常数。",
            ],
            "tables": [
                {
                    "headers": ["类别", "示例", "处理方式", "原因"],
                    "rows": [
                        ["控制器设计参数", "横向预瞄时间表节点、纵向比例积分增益、切换速度", "作为待整定参数参与梯度优化", "直接影响闭环跟踪误差和指令平滑性"],
                        ["车辆物理常数", "轴距、质量、轮胎半径、转向比、风阻/滚阻", "作为固定参数或训练域样本", "描述被控对象，不应在控制器整定中被误改"],
                        ["安全约束", "最大转角、最大加减速度、速率限制", "保留硬约束，训练时使用可传递梯度的替代处理", "确保整定结果不突破物理和安全边界"],
                        ["监控或辅助输出", "不参与反馈的近预瞄监控量", "不作为待整定参数", "变化不会有效降低闭环损失，纳入优化会增加噪声"],
                    ],
                }
            ],
        },
        {
            "level": 2,
            "title": "4.4 非光滑控制逻辑的可微处理",
            "paras": [
                "对于查找表，本方法保持横轴断点不变，只将纵轴节点值作为可调参数。线性插值天然可以把梯度分配给相邻节点。",
                "对于输出限幅和速率限制，本方法在前向计算中仍执行硬限幅，使控制指令满足工程边界；在反向传播中采用直通估计，使梯度能够传回限幅前的控制目标。",
                "对于依赖车速、曲率或加速度阈值的条件分支，本方法在训练模式下使用连续权重混合两个分支，在验证模式下恢复原始硬分支。",
                "对于最近轨迹点选择，本方法将离散索引选择与梯度链路隔离，避免最近点跳变带来错误梯度；对于远预瞄时间等可调时间参数，本方法使用按时间线性插值的轨迹查询，使预瞄时间变化能够影响损失。",
                "以上处理使控制器在训练中可回传梯度，同时避免把原工程控制器改写成完全不同的控制律。",
            ],
        },
        {
            "level": 2,
            "title": "4.5 批量并行和鲁棒训练",
            "paras": [
                "在一个实施例中，系统将8类轨迹和6个速度段展开为48条训练轨迹，并把每条轨迹按多个车辆参数域复制。若每轮采样K组车辆参数，则有效批量大小为48×K。所有批量元素在同一时间步同步执行轨迹查询、控制器计算、动力学积分和损失统计。",
                "车辆参数域可以包括牵引车质量、前轴侧偏刚度和后轴侧偏刚度，转动惯量可与质量联动变化，避免出现不符合物理规律的组合。反馈噪声可以作用在车辆位置、航向、速度和横摆角速度上，指令抖动可以作用在转向角和车轮扭矩上。验证阶段关闭这些扰动，使结果反映干净工况下的控制性能。",
            ],
        },
        {
            "level": 2,
            "title": "4.6 训练后验证与产物",
            "paras": [
                "训练完成后，系统自动保存整定参数配置，并在验证模式下运行调参前和调参后的对比仿真。验证覆盖训练轨迹同类的多速度场景，还可以加入园区综合路线等未参与训练的验证场景。验证输出包括轨迹跟踪图、横向误差图、速度误差图、转向角图、加速度图、训练摘要、参数变化图和实验日志。",
                "由于验证模式使用原始硬分支和硬限幅逻辑，因此该步骤能够发现只在平滑训练近似中有效、但在工程控制器中无效的参数，从而提高整定结果可部署性。",
            ],
        },
        {
            "level": 1,
            "title": "5. 附图及其说明",
            "paras": [
                "图1是车辆横纵向控制器可微闭环参数整定总体流程图。",
                "图2是可微训练路径与硬逻辑验证路径的关系示意图。",
                "图3是多轨迹多车辆参数域并行鲁棒整定流程图。",
                "图4是实施例一的训练损失曲线。",
                "图5是实施例一的控制器参数变化图。",
                "图6是实施例一的训练摘要和49场景硬逻辑验证统计图。",
                "图7是实施例一的调参前后轨迹跟踪对比图。",
                "图8是实施例一的调参前后横向误差对比图。",
            ],
        },
        {
            "level": 1,
            "title": "6. 具体实施例",
            "paras": [
                "以下实施例基于当前项目中的自动驾驶控制器复现和仿真训练结果。实施例仅用于说明本发明的可实施性，不限制本发明的保护范围。",
            ],
        },
        {
            "level": 2,
            "title": "6.1 实施例一：基于牵引车动力学和残差模型的闭环整定",
            "paras": [
                "在本实施例中，被控对象为牵引车动力学模型，车辆动力学采用机理模型和冻结的残差模型组合。控制器为横向多点预瞄控制器和纵向级联控制器，控制周期为50Hz，训练轨迹为8类轨迹乘以6个速度段，共48条训练轨迹。训练轮数为6，时间截断窗口为150步。",
                "训练结果显示，综合损失从3.9303下降到3.0535，下降22.31%。使用硬逻辑验证路径在49个场景中对比调参前后表现，其中38个场景横向误差均方根下降，43个场景航向误差均方根下降；横向误差变化的平均值为-11.20%，航向误差变化的平均值为-15.11%。",
                "参数变化显示，横向4组查找表均发生非零调整；纵向低速和高速速度环增益、站位环增益和切换速度也发生调整。该结果说明损失能够同时穿过横向预瞄、纵向速度环和车辆动力学链路回到控制器参数。",
            ],
        },
        {
            "level": 2,
            "title": "6.2 实施例二：车辆物理参数域随机化整定",
            "paras": [
                "在本实施例中，每轮训练采样4组牵引车质量、前轴侧偏刚度和后轴侧偏刚度，将48条轨迹复制到4个车辆域，共192条批量闭环仿真。牵引车质量采样范围为名义值的正负10%，前后轴侧偏刚度采样范围为名义值的正负20%。",
                "在纯机理模型实施例中，综合损失从2.7083下降到1.5639，下降42.26%。硬逻辑验证路径中，49个场景里有47个场景横向误差下降，47个场景航向误差下降；横向误差变化的平均值为-29.32%，航向误差变化的平均值为-24.87%。",
                "该实施例说明，通过在训练中暴露车辆物理参数的不确定性，控制器参数不是只适配单一名义车辆，而是在一族车辆参数上获得更稳定的跟踪表现。",
            ],
        },
        {
            "level": 2,
            "title": "6.3 实施例三：叠加反馈噪声和指令抖动的鲁棒整定",
            "paras": [
                "在本实施例中，车辆参数域随机化与反馈噪声、指令抖动同时启用。反馈噪声作用在位置、航向、速度和横摆角速度等控制器输入上，指令抖动作用在转向和扭矩执行指令上。",
                "训练结果显示，综合损失从4.6378下降到3.5931，下降22.53%。硬逻辑验证路径中，49个场景里有43个场景横向误差下降，33个场景航向误差下降；横向误差变化平均值为-17.61%。该结果说明，本发明能够把外部扰动纳入训练环境，同时仍通过无扰动硬逻辑路径评估整定效果。",
            ],
        },
        {
            "level": 1,
            "title": "7. 本发明的有益效果",
            "paras": [
                "第一，本发明将车辆横纵向控制器、车辆动力学和多场景评价指标统一到闭环时间链路中，使控制器参数能够根据跟踪误差自动获得调整方向，减少人工试错。",
                "第二，本发明通过双模式控制器兼顾可训练性和工程一致性。训练模式解决梯度回传问题，验证模式保留原始硬逻辑，从而降低训练近似与部署行为不一致的风险。",
                "第三，本发明直接整定具有明确工程含义的控制器参数，例如预瞄时间查找表和比例积分增益，整定结果可解释、可回写到配置文件，也便于工程师继续审查。",
                "第四，本发明支持多轨迹、多速度段和多车辆参数域并行训练，可以同时优化换道、弯道、S弯、弯前减速等不同工况下的整体表现。",
                "第五，本发明通过训练后自动验证和自动出图，形成从整定输入、参数变化、仿真指标到结果配置的完整证据链，便于参数版本管理和工程复现。",
            ],
        },
        {
            "level": 1,
            "title": "8. 希望保护的重点",
            "paras": [
                "1. 一种把车辆横纵向控制器、轨迹查询和车辆动力学模型展开为可自动求梯度的闭环时间链路，并用跟踪损失自动整定控制器参数的方法。",
                "2. 一种控制器双模式复现方法：同一控制器参数在训练模式下采用可微处理，在验证模式下采用原始硬分支、硬限幅和硬速率限制，从而保证整定结果可部署验证。",
                "3. 一种面向工程控制器非光滑逻辑的可微处理组合，包括查找表节点梯度分配、速率限制直通估计、条件分支连续混合、最近点选择隔离和可调时间预瞄插值。",
                "4. 一种横纵向控制器参数分类和投影方法，将预瞄时间、收敛时间、比例积分增益、速度切换点等作为可调参数，将车辆物理常数和安全约束作为固定参数或训练域样本，并在每轮更新后执行物理约束投影。",
                "5. 一种多轨迹、多速度段、多车辆参数域的批量并行闭环整定方法，将训练轨迹与采样车辆域展开为批量元素，同步推进仿真并对损失进行按轨迹归一化和按域统计。",
                "6. 一种训练后自动产物生成和硬逻辑复验方法，将整定参数导出为配置文件，并自动生成训练曲线、参数变化、调参前后轨迹和误差对比以及实验日志。",
                "7. 上述方法在机理车辆模型、机理与残差模型组合、牵引车与挂车双体动力学模型中的应用。",
            ],
        },
        {
            "level": 1,
            "title": "9. 可选专利布局建议",
            "paras": [
                "建议主案采用当前首选名称，保护“闭环可微仿真参数自动整定”的总方法。若后续需要扩大布局，可拆分出两个从属方向：其一是“控制器双模式可微复现及硬逻辑验证方法”，其二是“面向车辆动力学参数不确定性的鲁棒可微整定方法”。",
                "不建议主案只命名为“可微控制仿真环境的自动化搭建方法”。该名称容易把创新点落到环境搭建工具链上，而当前项目更强的技术贡献在于真实控制器参数如何在闭环车辆动力学中自动整定并完成硬逻辑验证。",
            ],
        },
        {
            "level": 1,
            "title": "10. 项目实现依据",
            "paras": [
                "控制器复现和参数分类可参考项目文件：sim/controller/lat_truck.py、sim/controller/lon.py、docs/tunable_params_analysis.md。",
                "闭环仿真和双路径验证可参考项目文件：sim/sim_loop.py、sim/common.py、sim/optim/post_training.py。",
                "批量并行训练、域随机化、反馈噪声和指令抖动可参考项目文件：sim/optim/train_batch.py、docs/plans/2026-05-08-domain-randomization-design.md、docs/plans/2026-05-08-aggressive-dr-noise-dither-design.md。",
                "实验结果图和日志可参考目录：sim/results/training/truck_trailer/20260526_123421_mlp0525、sim/results/training/truck_trailer/20260508_133208_nomlp_dr、sim/results/training/truck_trailer/20260509_181719_mlp0509_dr+noise+dither。",
            ],
        },
    ]


def md_from_sections(sections: list[dict]) -> str:
    out: list[str] = []
    for sec in sections:
        level = sec["level"]
        title = sec["title"]
        out.append(f"{'#' * (level + 1)} {title}\n")
        for para in sec.get("paras", []):
            if para.startswith(tuple(f"{i}." for i in range(1, 10))):
                out.append(f"{para}\n")
            else:
                out.append(f"{para}\n")
        for table in sec.get("tables", []):
            headers = table["headers"]
            out.append("| " + " | ".join(headers) + " |")
            out.append("| " + " | ".join(["---"] * len(headers)) + " |")
            for row in table["rows"]:
                out.append("| " + " | ".join(row) + " |")
            out.append("")
    out.append("## 附图\n")
    figs = [
        ("图1", "figures/fig1_overall_flow.png", "车辆横纵向控制器可微闭环参数整定总体流程"),
        ("图2", "figures/fig2_dual_path.png", "可微训练路径与硬逻辑验证路径"),
        ("图3", "figures/fig3_batched_domain_randomization.png", "多轨迹多域并行鲁棒整定流程"),
        ("图4", "figures/fig4_loss_curve.png", "实施例一训练损失曲线"),
        ("图5", "figures/fig5_parameter_changes.png", "实施例一控制器参数变化图"),
        ("图6", "figures/fig6_training_summary.png", "实施例一训练摘要和硬逻辑验证统计图"),
        ("图7", "figures/fig7_comparison_trajectory.png", "实施例一调参前后轨迹跟踪对比图"),
        ("图8", "figures/fig8_comparison_lateral_error.png", "实施例一调参前后横向误差对比图"),
    ]
    for label, path, caption in figs:
        out.append(f"![{label} {caption}]({path})\n")
    return "\n".join(out).strip() + "\n"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill.replace("#", ""))
    tc_pr.append(shd)


def set_run_font(run, size: int | None = None, east_asia: str = "宋体",
                 latin: str = "Calibri", bold: bool | None = None) -> None:
    if bold is not None:
        run.bold = bold
    if size is not None:
        run.font.size = Pt(size)
    run.font.name = latin
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) < 12 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, size=9, east_asia="宋体", bold=bold)


def apply_doc_styles(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after, east in [
        ("Heading 1", 16, "2E74B5", 16, 8, "黑体"),
        ("Heading 2", 13, "2E74B5", 12, 6, "黑体"),
        ("Heading 3", 12, "1F4D78", 8, 4, "黑体"),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), east)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    run = p.add_run(text)
    set_run_font(run, size=11, east_asia="宋体")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    for i, h in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], "F2F4F7")
        set_cell_text(table.rows[0].cells[i], h, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text)
    doc.add_paragraph()


def add_captioned_image(
    doc: Document,
    img_path: Path,
    caption: str,
    width_in: float = 6.3,
    max_height_in: float = 7.2,
) -> None:
    if not img_path.exists():
        return
    with Image.open(img_path) as img:
        image_width, image_height = img.size
    height_in = width_in * image_height / image_width
    if height_in > max_height_in:
        width_in = max_height_in * image_width / image_height

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_run_font(r, size=10, east_asia="宋体")


def build_docx(sections: list[dict]) -> None:
    doc = Document()
    apply_doc_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("专利技术交底书")
    set_run_font(run, size=22, east_asia="黑体", bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run(PREFERRED_TITLE)
    set_run_font(r, size=15, east_asia="黑体", bold=True)

    meta = [
        ["文档版本", "V1"],
        ["形成日期", "2026-07-09"],
        ["推荐主案名称", PREFERRED_TITLE],
        ["参考样例", "专利撰写目录中 leader 既有技术交底书组织方式"],
        ["项目依据", "differentiable_control / sim 可微控制参数整定项目"],
    ]
    add_table(doc, ["项目", "内容"], meta)

    for sec in sections[1:]:
        level = sec["level"]
        if level == 1:
            doc.add_heading(sec["title"], level=1)
        elif level == 2:
            doc.add_heading(sec["title"], level=2)
        else:
            doc.add_heading(sec["title"], level=3)
        for para in sec.get("paras", []):
            add_para(doc, para)
        for table in sec.get("tables", []):
            add_table(doc, table["headers"], table["rows"])

    doc.add_heading("附图", level=1)
    figs = [
        ("图1 车辆横纵向控制器可微闭环参数整定总体流程", "fig1_overall_flow.png", 6.3),
        ("图2 可微训练路径与硬逻辑验证路径", "fig2_dual_path.png", 6.3),
        ("图3 多轨迹多域并行鲁棒整定流程", "fig3_batched_domain_randomization.png", 6.3),
        ("图4 实施例一训练损失曲线", "fig4_loss_curve.png", 6.2),
        ("图5 实施例一控制器参数变化图", "fig5_parameter_changes.png", 6.2),
        ("图6 实施例一训练摘要和硬逻辑验证统计图", "fig6_training_summary.png", 6.2),
        ("图7 实施例一调参前后轨迹跟踪对比图", "fig7_comparison_trajectory.png", 6.2),
        ("图8 实施例一调参前后横向误差对比图", "fig8_comparison_lateral_error.png", 6.2),
    ]
    for caption, name, width in figs:
        add_captioned_image(doc, FIG_DIR / name, caption, width)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = footer.add_run("专利技术交底书 V1")
    set_run_font(rr, size=9, east_asia="宋体")
    doc.save(DOCX_PATH)


def main() -> None:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    generate_flow_figures()
    copy_project_figures()
    sections = disclosure_sections()
    MD_PATH.write_text(md_from_sections(sections), encoding="utf-8")
    build_docx(sections)
    print(f"wrote {MD_PATH}")
    print(f"wrote {DOCX_PATH}")


if __name__ == "__main__":
    main()
